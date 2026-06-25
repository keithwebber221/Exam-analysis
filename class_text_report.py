"""
class_text_report.py
全班試卷分析文字報告模組
輸入：df, max_scores, paper_map, class_info, exam_info, absent_set,
       item_df, group_df, student_df, stats_df
輸出：Word (.docx) BytesIO
整合：由 app.py 呼叫 generate_text_report()
"""

import io, re
import numpy as np
import pandas as pd
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════
# 樣式工具
# ══════════════════════════════════════════════════════════════

def _set_font(run, size, bold=False, color=None):
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _no_page_break(p):
    pPr = p._p.get_or_add_pPr()
    for tag in ["w:pageBreakBefore", "w:keepNext", "w:keepLines"]:
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.remove(el)
        el = OxmlElement(tag)
        el.set(qn("w:val"), "0")
        pPr.append(el)


def _add_heading(doc, text, level=1, color=(0, 0, 0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    p.paragraph_format.space_after  = Pt(2)
    _no_page_break(p)
    r = p.add_run(text)
    _set_font(r, 13 if level == 1 else 11, bold=True, color=color)
    return p


def _add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    _no_page_break(p)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    _set_font(r, 10.5)
    return p


def _add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.6)
    _no_page_break(p)
    r = p.add_run(text)
    _set_font(r, 10.5, color=color)
    return p


def _add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    _no_page_break(p)
    r = p.add_run("─" * 60)
    _set_font(r, 8, color=(190, 190, 190))


def _add_spacer(doc, after_pt=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    _no_page_break(p)


def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_note_box(doc, label, body_lines,
                  bg_hex="EEF4FF", label_rgb=(31, 71, 136), body_rgb=(40, 40, 80)):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, bg_hex)

    trPr      = tbl.rows[0]._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    cantSplit.set(qn("w:val"), "0")
    trPr.append(cantSplit)

    cell.paragraphs[0].clear()
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(3)
    p0.paragraph_format.space_after  = Pt(1)
    _no_page_break(p0)
    rl = p0.add_run(f"{label}　")
    _set_font(rl, 10, bold=True, color=label_rgb)
    rb = p0.add_run(body_lines[0])
    _set_font(rb, 10, color=body_rgb)

    for line in body_lines[1:]:
        pe = cell.add_paragraph()
        pe.paragraph_format.space_before = Pt(0)
        pe.paragraph_format.space_after  = Pt(1)
        pe.paragraph_format.left_indent  = Cm(0.2)
        _no_page_break(pe)
        re_ = pe.add_run(line)
        _set_font(re_, 10, color=body_rgb)
    return tbl


# ══════════════════════════════════════════════════════════════
# 數據計算輔助
# ══════════════════════════════════════════════════════════════

def _get_parent_q(q_name):
    parts = q_name.split("_", 1)
    if len(parts) < 2:
        # No paper prefix — single-paper mode
        item = q_name
        paper = "P1"
        if "Part" in item or "part" in item:
            return paper, "PartA（選擇題）"
        # Match "Q1a" → "Q1",  "1a" → "1",  "Q10b" → "Q10",  "10b" → "10"
        m = re.match(r"(Q?[\d/]+)", item)
        if m:
            return paper, m.group(1)
        return paper, item
    paper, item = parts
    if "Part" in item or "part" in item:
        return paper, f"{paper}_PartA（選擇題）"
    m = re.match(r"(Q?[\d/]+)", item)
    if m:
        return paper, f"{paper}_{m.group(1)}"
    return paper, f"{paper}_{item}"


def _build_group_stats(df, max_scores, paper_map, absent_set, total_scores):
    """Build parent-question group statistics."""
    present_mask   = ~df.index.isin(absent_set)
    present_scores = df[present_mask]
    overall_rate   = total_scores.mean() / max_scores.sum() * 100

    parent_groups = defaultdict(list)
    for q in df.columns:
        _, parent = _get_parent_q(q)
        parent_groups[parent].append(q)

    rows = []
    for parent, qs in parent_groups.items():
        grp_max    = sum(max_scores[q] for q in qs)
        grp_scores = present_scores[qs].sum(axis=1)
        grp_avg    = grp_scores.mean()
        grp_rate   = grp_avg / grp_max * 100 if grp_max else 0

        diff = ("容易" if grp_rate >= 70 else "困難" if grp_rate < 40 else "適中")
        # Clean up display name
        display = parent
        for prefix, label in [("P1_", ""), ("P2_", "卷二 "), ("P3_", "卷三 "), ("P4_", "卷四 ")]:
            display = display.replace(prefix, label)
        display = display.strip()
        rows.append({
            "parent": parent, "display": display,
            "qs": qs, "max": grp_max, "avg": round(grp_avg, 2),
            "rate": round(grp_rate, 1), "diff": diff,
            "paper": qs[0].split("_")[0] if "_" in qs[0] else "P1",
        })

    group_df2 = pd.DataFrame(rows).sort_values("rate", ascending=False)
    return group_df2, overall_rate


# ══════════════════════════════════════════════════════════════
# 各部分建構函式
# ══════════════════════════════════════════════════════════════

def _build_cover(doc, exam_info):
    subject_label   = exam_info.get("subject_label", "")
    form_label      = exam_info.get("form_label", "")
    year_label      = exam_info.get("year_label", "")
    exam_type_label = exam_info.get("exam_type_label", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    _no_page_break(p)
    r = p.add_run("全班試卷分析文字報告")
    _set_font(r, 20, bold=True, color=(0, 0, 0))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    _no_page_break(p2)
    r2 = p2.add_run(f"{year_label}　{exam_type_label}　{form_label}　{subject_label}")
    _set_font(r2, 11, color=(0, 0, 0))

    _add_spacer(doc, 2)
    _add_divider(doc)

    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_before = Pt(2)
    toc_p.paragraph_format.space_after  = Pt(2)
    _no_page_break(toc_p)
    r_toc = toc_p.add_run(
        "本報告共分五部分：一、試卷總覽　二、各題表現分析　三、大題分析　四、學生分層分析　五、後續跟進建議"
    )
    _set_font(r_toc, 9.5, color=(0, 0, 0))

    _add_spacer(doc, 2)
    _add_note_box(
        doc, label="📌 名詞說明",
        bg_hex="EEF4FF", label_rgb=(31, 71, 136), body_rgb=(40, 40, 80),
        body_lines=[
            "閱讀本報告前，請留意以下兩個常用指標的含義：",
            "",
            "【pp（百分點 / percentage points）】兩個百分比之間的直接差距。"
            "例如得分率 34% 與全級平均 61%，相差 27pp。▲ 高於全級平均；▼ 低於全級平均。",
            "",
            "【鑑別度（Discrimination Index）】衡量題目能否有效區分高、低能力學生，數值介乎 −1 至 1，數值愈高愈理想。",
            "    • 0.40 或以上 → 優良　• 0.30–0.39 → 良好　"
            "• 0.20–0.29 → 尚可（可考慮修訂）　• 0.20 以下 → 不佳（建議重新設計）",
        ]
    )
    _add_spacer(doc, 2)


def _build_overview(doc, exam_info, n_total, n_absent, n_present,
                    n_pass, pass_pct, mean_score, mean_pct,
                    max_score_val, min_score_val, std_score, median_score,
                    total_max, pass_threshold, paper_max_str):
    subject_label   = exam_info.get("subject_label", "")
    form_label      = exam_info.get("form_label", "")
    year_label      = exam_info.get("year_label", "")
    exam_type_label = exam_info.get("exam_type_label", "")
    pass_rate_pct   = int(exam_info.get("pass_rate", 0.4) * 100)

    _add_heading(doc, "一、試卷總覽")

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    rows_data = [
        ["考試名稱", f"{year_label} {exam_type_label}", "科目", subject_label],
        ["年級",     form_label, "總滿分", f"{total_max} 分{paper_max_str}"],
        ["及格線",   f"{int(pass_threshold)} 分（{pass_rate_pct}%）",
         "出席人數", f"{n_present} 人（全班 {n_total} 人，缺席 {n_absent} 人）"],
    ]
    for ri, row_data in enumerate(rows_data):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.text = val
            run  = cell.paragraphs[0].runs[0]
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(1)
            _no_page_break(cell.paragraphs[0])
            _set_font(run, 9.5, bold=(ci % 2 == 0),
                      color=(0, 0, 0))

    _add_spacer(doc, 1)

    overall_comment = (
        f"整體表現理想，及格率達 {pass_pct:.1f}%，反映大部分同學已掌握本科核心概念。"
        if pass_pct >= 85 else
        f"整體表現尚可，及格率為 {pass_pct:.1f}%，仍有部分同學需要針對性跟進。"
        if pass_pct >= 70 else
        f"整體表現欠佳，及格率僅 {pass_pct:.1f}%，建議全面檢視教學重點。"
    )
    _add_body(doc,
        f"全班共 {n_present} 名同學出席考試，及格人數為 {n_pass} 人，"
        f"及格率為 {pass_pct:.1f}%。{overall_comment}"
    )
    _add_body(doc,
        f"全班平均分為 {mean_score:.1f} 分（得分率 {mean_pct:.1f}%），"
        f"最高分 {max_score_val:.0f} 分，最低分 {min_score_val:.0f} 分，"
        f"標準差 {std_score:.1f}，中位數 {median_score:.1f} 分。"
    )
    _add_divider(doc)


def _build_item_analysis(doc, item_df, overall_rate):
    _add_heading(doc, "二、各題表現分析")
    _add_body(doc, "以下按各題得分率及鑑別度，分類呈現表現情況：")

    # 難度欄位名稱可能不同，統一計算
    if "得分率%" in item_df.columns:
        rate_col = "得分率%"
    else:
        item_df = item_df.copy()
        item_df["得分率%"] = (item_df["平均分"] / item_df["滿分"] * 100).round(1)
        rate_col = "得分率%"

    weak_qs   = item_df[item_df[rate_col] < 40].sort_values(rate_col)
    strong_qs = item_df[item_df[rate_col] >= 70].sort_values(rate_col, ascending=False)
    poor_disc = item_df[item_df["鑑別度"] < 0.2] if "鑑別度" in item_df.columns else pd.DataFrame()

    # 難度標籤（若無難度欄位則自行計算）
    def _diff(rate):
        return "容易" if rate >= 70 else "困難" if rate < 40 else "適中"

    def _disc_label(d):
        if d >= 0.4: return "優良"
        if d >= 0.3: return "良好"
        if d >= 0.2: return "尚可"
        return "不佳"

    q_col = "題目" if "題目" in item_df.columns else item_df.columns[0]

    # 甲：弱項
    _add_heading(doc, "（甲）表現較弱的題目（得分率低於 40%）", level=2, color=(180, 40, 40))
    if len(weak_qs) > 0:
        for _, row in weak_qs.iterrows():
            qn  = str(row[q_col]).replace("P1_", "卷一 ").replace("P2_", "卷二 ").replace("P3_", "卷三 ").replace("P4_", "卷四 ")
            d   = row.get("鑑別度", 0)
            dl  = _disc_label(d)
            dif = _diff(row[rate_col])
            _add_bullet(doc,
                f"{qn}（滿分 {int(row['滿分'])} 分）：平均 {row['平均分']:.1f} 分，"
                f"得分率 {row[rate_col]:.1f}%，屬{dif}題，鑑別度 {d:.2f}（{dl}）。建議補充相關概念。"
            )
    else:
        _add_body(doc, "本次考試所有題目得分率均達 40% 以上。", indent=True)

    # 乙：強項
    _add_heading(doc, "（乙）表現較佳的題目（得分率達 70% 或以上）", level=2, color=(0, 120, 60))
    for _, row in strong_qs.head(5).iterrows():
        qn = str(row[q_col]).replace("P1_", "卷一 ").replace("P2_", "卷二 ").replace("P3_", "卷三 ").replace("P4_", "卷四 ")
        _add_bullet(doc,
            f"{qn}（滿分 {int(row['滿分'])} 分）：平均 {row['平均分']:.1f} 分，"
            f"得分率 {row[rate_col]:.1f}%，同學整體掌握良好。"
        )

    # 丙：鑑別度不佳
    _add_heading(doc, "（丙）鑑別度不佳的題目（鑑別度低於 0.2）", level=2, color=(160, 90, 0))
    if len(poor_disc) > 0:
        for _, row in poor_disc.iterrows():
            qn = str(row[q_col]).replace("P1_", "卷一 ").replace("P2_", "卷二 ").replace("P3_", "卷三 ").replace("P4_", "卷四 ")
            _add_bullet(doc,
                f"{qn}：鑑別度 {row['鑑別度']:.2f}，未能有效區分不同能力的學生，建議檢視題目設計。"
            )
    else:
        _add_body(doc, "本次考試各題鑑別度整體理想。", indent=True)

    _add_divider(doc)


def _build_group_analysis(doc, df, max_scores, paper_map, absent_set,
                           total_scores, item_df):
    group_df2, overall_rate = _build_group_stats(
        df, max_scores, paper_map, absent_set, total_scores)

    _add_heading(doc, "三、大題分析")
    _add_body(doc, f"以下按大題匯總各小題表現，並與全卷整體得分率（{overall_rate:.1f}%）作比較：")

    q_col     = "題目" if "題目" in item_df.columns else item_df.columns[0]
    rate_col  = "得分率%" if "得分率%" in item_df.columns else None
    papers_in = sorted(group_df2["paper"].unique())

    multi_paper_mode = len(papers_in) > 1

    for paper_key in papers_in:
        label_map = {"P1": "卷一（Paper 1）", "P2": "卷二（Paper 2）",
                     "P3": "卷三（Paper 3）", "P4": "卷四（Paper 4）"}
        if multi_paper_mode:
            paper_label = label_map.get(paper_key, paper_key)
            _add_heading(doc, f"【{paper_label}】", level=2, color=(0, 0, 0))

        paper_rows = group_df2[group_df2["paper"] == paper_key].sort_values("parent").to_dict("records")

        for row in paper_rows:
            rate    = row["rate"];  delta   = rate - overall_rate
            mx      = row["max"];   avg     = row["avg"]
            diff    = row["diff"];  qs      = row["qs"];  display = row["display"]

            rate_color   = (0, 120, 60) if rate >= 70 else (160, 100, 0) if rate >= 40 else (180, 40, 40)
            delta_symbol = "▲" if delta >= 0 else "▼"
            delta_color  = (0, 120, 60) if delta >= 0 else (180, 40, 40)

            # sub-question breakdown
            sub_parts = []
            for q in qs:
                q_short  = q.split("_", 1)[1] if "_" in q else q
                q_row    = item_df[item_df[q_col] == q]
                if len(q_row):
                    if rate_col:
                        q_rate = q_row[rate_col].values[0]
                    else:
                        q_rate = q_row["平均分"].values[0] / q_row["滿分"].values[0] * 100
                    q_delta = q_rate - overall_rate
                    sym = "▲" if q_delta >= 0 else "▼"
                    sub_parts.append(f"{q_short} {q_rate:.0f}%（{sym}{abs(q_delta):.0f}pp）")

            # comment
            if   delta >= 15:  comment = f"表現顯著高於全卷平均（+{delta:.1f}pp），屬本次考試強項大題。"
            elif delta >= 5:   comment = f"表現略高於全卷平均（+{delta:.1f}pp），同學整體掌握尚可。"
            elif delta >= -5:  comment = f"表現與全卷平均相近（{delta:+.1f}pp），屬正常水平。"
            elif delta >= -15: comment = f"表現低於全卷平均（{delta:.1f}pp），建議加強相關概念的練習。"
            else:              comment = f"表現顯著低於全卷平均（{delta:.1f}pp），屬本次最弱大題，建議優先補底。"

            # worst sub-question
            sub_rates_list = []
            for q in qs:
                q_row = item_df[item_df[q_col] == q]
                if len(q_row):
                    r_val = q_row[rate_col].values[0] if rate_col else q_row["平均分"].values[0] / q_row["滿分"].values[0] * 100
                    sub_rates_list.append((q, r_val))
            weakness_str = ""
            if len(sub_rates_list) > 1:
                worst_q, worst_r = min(sub_rates_list, key=lambda x: x[1])
                wq_short = worst_q.split("_", 1)[1] if "_" in worst_q else worst_q
                weakness_str = f" 其中 {wq_short} 得分率 {worst_r:.0f}%，為本大題最弱小題。"

            # Line 1: title + stats + delta
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Cm(0.3)
            _no_page_break(p)
            r1 = p.add_run(f"{display}　");                _set_font(r1, 10.5, bold=True, color=rate_color)
            r2 = p.add_run(f"滿分 {int(mx)} 分　平均 {avg:.1f} 分　得分率 {rate:.1f}%（{diff}）　"); _set_font(r2, 10.5, color=(0, 0, 0))
            r3 = p.add_run(f"{delta_symbol} {abs(delta):.1f}pp（{"高於" if delta >= 0 else "低於"}全級平均 {overall_rate:.1f}%）"); _set_font(r3, 10.5, bold=True, color=delta_color)

            # Line 2: sub-question
            if sub_parts:
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.left_indent  = Cm(1.0)
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after  = Pt(0)
                _no_page_break(p_sub)
                r_sub = p_sub.add_run(f"各小題得分率：{"　".join(sub_parts)}")
                _set_font(r_sub, 9.5, color=(0, 0, 0))

            # Line 3: comment
            p_comm = doc.add_paragraph()
            p_comm.paragraph_format.left_indent  = Cm(1.0)
            p_comm.paragraph_format.space_before = Pt(0)
            p_comm.paragraph_format.space_after  = Pt(5)
            _no_page_break(p_comm)
            r_comm = p_comm.add_run(f"▸ {comment}{weakness_str}")
            _set_font(r_comm, 9.5, color=(0, 0, 0))

    best_grp  = group_df2.iloc[0]
    worst_grp = group_df2.iloc[-1]
    _add_body(doc,
        f"綜合而言，以「{best_grp['display']}」表現最佳"
        f"（{best_grp['rate']:.1f}%，▲{best_grp['rate'] - overall_rate:.1f}pp），"
        f"以「{worst_grp['display']}」表現最遜"
        f"（{worst_grp['rate']:.1f}%，▼{overall_rate - worst_grp['rate']:.1f}pp），"
        f"建議列為優先補底重點。"
    )
    _add_divider(doc)
    return group_df2, overall_rate, best_grp, worst_grp


def _build_tier_analysis(doc, total_scores, total_max, pass_threshold, n_present):
    _add_heading(doc, "四、學生分層分析")

    tier_top = total_scores[total_scores >= total_max * 0.8]
    tier_mid = total_scores[(total_scores >= pass_threshold) & (total_scores < total_max * 0.8)]
    tier_low = total_scores[total_scores < pass_threshold]

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"
    cm_list = [(31, 71, 136), (0, 120, 60), (160, 100, 0), (180, 40, 40)]
    for ri, row_data in enumerate([
        ["分層", "標準", "人數", "佔出席比例"],
        ["優秀",   f"≥ {int(total_max * 0.8)} 分（80%+）",                  str(len(tier_top)), f"{len(tier_top) / n_present * 100:.1f}%"],
        ["合格",   f"{int(pass_threshold)}–{int(total_max * 0.8) - 1} 分",  str(len(tier_mid)), f"{len(tier_mid) / n_present * 100:.1f}%"],
        ["待改善", f"< {int(pass_threshold)} 分",                            str(len(tier_low)), f"{len(tier_low) / n_present * 100:.1f}%"],
    ]):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.text = val
            run  = cell.paragraphs[0].runs[0]
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(1)
            _no_page_break(cell.paragraphs[0])
            _set_font(run, 9.5, bold=(ri == 0 or ci == 0),
                      color=cm_list[ri] if (ci == 0 or ri == 0) else (30, 30, 30))

    _add_spacer(doc, 1)

    if len(tier_low) > 0:
        _add_body(doc,
            f"以下 {len(tier_low)} 名同學未達及格線，建議優先安排個別跟進："
            f"{"、".join(tier_low.index.tolist())}。"
        )
    if len(tier_top) > 0:
        _add_body(doc,
            f"以下 {len(tier_top)} 名同學表現優秀，可提供延伸練習："
            f"{"、".join(tier_top.index.tolist())}。"
        )
    _add_divider(doc)
    return tier_top, tier_mid, tier_low


def _build_suggestions(doc, weak_qs, poor_disc, worst_grp, overall_rate,
                        tier_low, std_score, pass_pct, item_df):
    _add_heading(doc, "五、後續跟進建議")

    q_col    = "題目" if "題目" in item_df.columns else item_df.columns[0]
    rate_col = "得分率%" if "得分率%" in item_df.columns else None

    if rate_col is None:
        weak_qs = item_df.copy()
        weak_qs["得分率%"] = (weak_qs["平均分"] / weak_qs["滿分"] * 100).round(1)
        weak_qs = weak_qs[weak_qs["得分率%"] < 40].sort_values("得分率%")
        rate_col = "得分率%"

    suggestions = []

    if len(weak_qs) > 0:
        wq_names = "、".join([
            str(r).replace("P1_", "卷一 ").replace("P2_", "卷二 ")
                  .replace("P3_", "卷三 ").replace("P4_", "卷四 ")
            for r in weak_qs[q_col].tolist()
        ])
        suggestions.append(
            f"【重點補底】{wq_names} 等題目得分率低於 40%，"
            f"建議下次課堂重點複習，並提供針對性練習。"
        )

    suggestions.append(
        f"【大題跟進】「{worst_grp['display']}」為得分率最低大題"
        f"（{worst_grp['rate']:.1f}%，▼{overall_rate - worst_grp['rate']:.1f}pp），"
        f"建議下一教學單元前安排專題複習。"
    )

    if len(tier_low) > 0:
        suggestions.append(
            f"【個別跟進】共 {len(tier_low)} 名同學未達及格線，"
            f"建議安排個別面談及補底支援。"
        )

    if "鑑別度" in item_df.columns:
        poor_disc_df = item_df[item_df["鑑別度"] < 0.2]
        if len(poor_disc_df) > 0:
            pd_names = "、".join([
                str(r).replace("P1_", "卷一 ").replace("P2_", "卷二 ")
                      .replace("P3_", "卷三 ").replace("P4_", "卷四 ")
                for r in poor_disc_df[q_col].tolist()
            ])
            suggestions.append(
                f"【題目檢視】{pd_names} 鑑別度低於 0.2，"
                f"建議下次命題時重新設計以提升評估效度。"
            )

    if std_score > 20:
        suggestions.append(
            f"【差異照顧】全班標準差 {std_score:.1f}，能力差距較大，"
            f"建議課堂加入分層練習，照顧不同程度學生。"
        )

    suggestions.append(
        f"【持續監察】建議下次考試後比較大題得分率趨勢，"
        f"追蹤各弱項大題的改善情況。"
    )

    for s in suggestions:
        _add_bullet(doc, s)

    _add_spacer(doc, 2)
    _add_divider(doc)

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.paragraph_format.space_before = Pt(1)
    _no_page_break(p_footer)
    r_footer = p_footer.add_run("本報告由試卷分析系統自動生成　如有疑問請核對原始數據")
    _set_font(r_footer, 8.5, color=(160, 160, 160))


# ══════════════════════════════════════════════════════════════
# 主入口
# ══════════════════════════════════════════════════════════════

def generate_text_report(
    df, max_scores, paper_map, class_info, exam_info,
    absent_set, item_df, group_df, student_df, stats_df
) -> bytes:
    """
    生成全班試卷分析文字報告（Word .docx）。

    Parameters
    ----------
    df          : pd.DataFrame  學生得分矩陣（index=姓名，columns=題目）
    max_scores  : pd.Series     各題滿分
    paper_map   : dict          題目 -> 試卷編號
    class_info  : pd.DataFrame  班別名冊（含 班別/班號/中文姓名）
    exam_info   : dict          考試資訊（subject_label, form_label, year_label, ...）
    absent_set  : set           缺席學生姓名集合
    item_df     : pd.DataFrame  試題分析表（含 題目/滿分/平均分/鑑別度 等欄位）
    group_df    : pd.DataFrame  大題分析表（來自 ea.question_group_analysis）
    student_df  : pd.DataFrame  學生成績表
    stats_df    : pd.DataFrame  全班統計表

    Returns
    -------
    bytes  Word docx 內容
    """
    # ── 基本統計 ──
    present_mask   = ~df.index.isin(absent_set)
    present_scores = df[present_mask]
    total_max      = int(max_scores.sum())
    pass_rate      = float(exam_info.get("pass_rate", 0.4))
    pass_threshold = total_max * pass_rate
    total_scores   = present_scores.sum(axis=1)

    n_total   = len(df)
    n_absent  = len(absent_set)
    n_present = n_total - n_absent
    n_pass    = int((total_scores >= pass_threshold).sum())
    pass_pct  = n_pass / n_present * 100 if n_present else 0

    mean_score    = float(total_scores.mean())
    mean_pct      = mean_score / total_max * 100
    max_score_val = float(total_scores.max())
    min_score_val = float(total_scores.min())
    std_score     = float(total_scores.std())
    median_score  = float(total_scores.median())

    # ── 試卷最大分拆解字串 ──
    papers = sorted(set(paper_map.values()))
    if len(papers) > 1:
        parts = []
        for p_key in papers:
            qs_in_p = [q for q, pv in paper_map.items() if pv == p_key]
            p_max   = int(sum(max_scores[q] for q in qs_in_p))
            label   = p_key.replace("P1", "卷一").replace("P2", "卷二").replace("P3", "卷三").replace("P4", "卷四")
            parts.append(f"{label} {p_max} 分")
        paper_max_str = "（" + "　".join(parts) + "）"
    else:
        paper_max_str = ""

    # ── 弱項 / 強項（item_df 增補得分率欄）──
    _item = item_df.copy()
    q_col = "題目" if "題目" in _item.columns else _item.columns[0]
    if "得分率%" not in _item.columns:
        _item["得分率%"] = (_item["平均分"] / _item["滿分"] * 100).round(1)
    weak_qs   = _item[_item["得分率%"] < 40].sort_values("得分率%")
    poor_disc = _item[_item["鑑別度"] < 0.2] if "鑑別度" in _item.columns else pd.DataFrame()

    # ── 建立 Word 文件 ──
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    _build_cover(doc, exam_info)
    _build_overview(
        doc, exam_info,
        n_total, n_absent, n_present, n_pass, pass_pct,
        mean_score, mean_pct, max_score_val, min_score_val,
        std_score, median_score, total_max, pass_threshold, paper_max_str
    )
    _build_item_analysis(doc, _item, mean_pct)
    _, overall_rate, best_grp, worst_grp = _build_group_analysis(
        doc, df, max_scores, paper_map, absent_set, total_scores, _item)
    tier_top, tier_mid, tier_low = _build_tier_analysis(
        doc, total_scores, total_max, pass_threshold, n_present)
    _build_suggestions(
        doc, weak_qs, poor_disc, worst_grp, overall_rate,
        tier_low, std_score, pass_pct, _item)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
