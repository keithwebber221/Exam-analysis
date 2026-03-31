#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
performance_tracker.py  ──  DSE 試卷分析系統 第二階段
學生成績跨試追蹤與比較系統
版本：v1.0  日期：2026-03-27
"""

import os, glob, re, ast
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import font_manager
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════
# 0. 中文字體設定
# ══════════════════════════════════════════════════════════════
def setup_chinese_font():
    """設定 matplotlib 中文字型，支援 macOS / Windows / Linux"""
    candidates = [
        # macOS
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        # Windows
        "C:/Windows/Fonts/msjh.ttc",
        "C:/Windows/Fonts/mingliu.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        # Linux
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                font_manager.fontManager.addfont(path)
                prop = font_manager.FontProperties(fname=path)
                name = prop.get_name()
                # 正確做法：family=sans-serif，再把中文字型插到最前
                plt.rcParams['font.family']     = 'sans-serif'
                plt.rcParams['font.sans-serif'] = [name] + plt.rcParams.get('font.sans-serif', [])
                plt.rcParams['axes.unicode_minus'] = False
                return name
            except Exception:
                continue
    plt.rcParams['font.family']        = 'sans-serif'
    plt.rcParams['axes.unicode_minus'] = False
    return 'DejaVu Sans'

FONT_NAME = setup_chinese_font()

# ══════════════════════════════════════════════════════════════
# 1. 掃描並讀取所有 analysis.xlsx
# ══════════════════════════════════════════════════════════════
EXAM_ORDER = {'T1T': 1, 'T1E': 2, 'T2T': 3, 'T2E': 4}
EXAM_LABELS = {
    'T1T': '上學期測驗', 'T1E': '上學期考試',
    'T2T': '下學期測驗', 'T2E': '下學期考試'
}

def parse_filename(fname):
    """
    從檔名提取年度、考試類型、年級、科目，返回4個值

    支援格式：
      新格式：2526_T1E_F5_BAFS_analysis.xlsx
      新格式（無科目）：2526_T1E_F5_analysis.xlsx
      舊格式：2526_T1E_analysis.xlsx

    Returns: (year, exam_type, form, subject)
             year/exam_type 為 None 表示無法解析
    """
    base = os.path.basename(fname)

    # 新格式（含年級+科目）：YYYY_TYPE_Fx_SUBJECT_analysis.xlsx
    m = re.match(
        r'(\d{4})_(T1T|T1E|T2T|T2E)_(F[1-6])_([^_]+)_analysis\.xlsx',
        base, re.IGNORECASE)
    if m:
        return m.group(1), m.group(2).upper(), m.group(3).upper(), m.group(4)

    # 新格式（含年級、無科目）：YYYY_TYPE_Fx_analysis.xlsx
    m = re.match(
        r'(\d{4})_(T1T|T1E|T2T|T2E)_(F[1-6])_analysis\.xlsx',
        base, re.IGNORECASE)
    if m:
        return m.group(1), m.group(2).upper(), m.group(3).upper(), ""

    # 舊格式：YYYY_TYPE_analysis.xlsx
    m = re.match(
        r'(\d{4})_(T1T|T1E|T2T|T2E)_analysis\.xlsx',
        base, re.IGNORECASE)
    if m:
        return m.group(1), m.group(2).upper(), "", ""

    return None, None, None, None

def sort_key(exam_tuple):
    """排序鍵：(年度, 考試類型) → 數字"""
    year, etype = exam_tuple
    return (int(year), EXAM_ORDER.get(etype, 99))

def scan_analysis_files(folder=".", filter_form=None, filter_subject=None):
    """
    掃描資料夾中所有 *_analysis.xlsx 檔案

    Parameters:
        filter_form    : 只掃描指定年級（如 "F5"），None = 全部
        filter_subject : 只掃描指定科目（如 "BAFS"），None = 全部
    """
    pattern = os.path.join(folder, "*_analysis.xlsx")
    files = glob.glob(pattern)
    results = []
    for f in files:
        year, etype, form, subject = parse_filename(f)
        if year and etype:
            # 過濾條件
            if filter_form and (form or "").upper() != filter_form.upper():
                continue
            if filter_subject and (subject or "").upper() != filter_subject.upper():
                continue
            results.append({
                'file':    f,
                'year':    year,
                'type':    etype,
                'form':    form,
                'subject': subject,
            })
    results.sort(key=lambda x: (x['form'], sort_key((x['year'], x['type']))))
    return results

def read_student_scores(filepath):
    """
    讀取 analysis.xlsx 的學生成績工作表
    使用模糊匹配工作表名稱（含「學生成績」關鍵字），
    避免因 emoji 差異導致找不到工作表
    """
    try:
        # 取得所有工作表名稱
        xl = pd.ExcelFile(filepath)
        sheet_names = xl.sheet_names

        # 模糊匹配：尋找包含「學生成績」的工作表
        target_sheet = None
        for s in sheet_names:
            if "學生成績" in s:
                target_sheet = s
                break

        if target_sheet is None:
            print(f"  ⚠️  {os.path.basename(filepath)} 找不到學生成績工作表")
            print(f"       可用工作表：{sheet_names}")
            return None

        df = pd.read_excel(filepath, sheet_name=target_sheet)

        # 欄位：姓名, 總分, 百分比(%), 排名
        # 支援新版（姓名）和舊版（student）兩種欄名
        name_col = None
        for candidate in ["姓名", "student"]:
            if candidate in df.columns:
                name_col = candidate
                break

        if name_col and "百分比(%)" in df.columns:
            df["百分比(%)"] = pd.to_numeric(
                df["百分比(%)"].astype(str).str.replace("%", "", regex=False),
                errors="coerce"
            )
            cols = [name_col, "總分", "百分比(%)", "排名"]
            # 同時讀取班別/班號（如存在）
            for extra in ["班別", "班號"]:
                if extra in df.columns:
                    cols.append(extra)
            result = df[cols].dropna(subset=[name_col])
            result = result.rename(columns={name_col: "姓名"})
            return result
        else:
            print(f"  ⚠️  {os.path.basename(filepath)} 工作表欄位不符，欄位：{list(df.columns)}")
            return None

    except Exception as e:
        print(f"  ⚠️  讀取 {os.path.basename(filepath)} 失敗：{e}")
    return None

# ══════════════════════════════════════════════════════════════
# 2. 建立跨試成績矩陣
# ══════════════════════════════════════════════════════════════
def build_tracking_matrix(exam_files, class_info_df=None):
    """
    建立學生 × 考試的成績追蹤矩陣

    Returns:
        pct_matrix : DataFrame  (學生 × 考試, 值為得分率或 NaN)
        rank_matrix: DataFrame  (學生 × 考試, 值為排名或 NaN)
        exam_labels: list of str (考試標籤)
        student_info: DataFrame (姓名, 班別, 班號)
    """
    if not exam_files:
        print("❌ 未找到任何 *_analysis.xlsx 檔案")
        return None, None, [], None

    # 收集所有學生、考試和班別班號
    all_students = set()
    exam_data    = {}   # key: "YYYY_TYPE"  val: {姓名: (百分比, 排名)}
    class_from_excel = {}  # 從 analysis.xlsx 收集的班別班號 {姓名: (班別, 班號)}

    for ef in exam_files:
        label = f"{ef['year']}_{ef['type']}"
        df = read_student_scores(ef['file'])
        if df is not None:
            exam_data[label] = dict(
                zip(df['姓名'], zip(df['百分比(%)'], df['排名']))
            )
            all_students.update(df['姓名'].tolist())
            # 收集班別/班號（每次考試都更新，以最新為準）
            if '班別' in df.columns and '班號' in df.columns:
                for _, row in df.iterrows():
                    name = row['姓名']
                    ban  = str(row.get('班別', '')).strip()
                    num  = str(row.get('班號', '')).strip()
                    if ban or num:
                        class_from_excel[name] = (ban, num)
            print(f"  ✅ 讀取 {os.path.basename(ef['file'])}：{len(df)} 位學生")

    if not exam_data:
        return None, None, [], None

    exam_labels_sorted = sorted(exam_data.keys(),
                                key=lambda x: sort_key(tuple(x.split('_'))))
    all_students = sorted(all_students)

    # 建立矩陣
    pct_data  = {lbl: {} for lbl in exam_labels_sorted}
    rank_data = {lbl: {} for lbl in exam_labels_sorted}

    for lbl in exam_labels_sorted:
        for student in all_students:
            if student in exam_data[lbl]:
                pct, rank = exam_data[lbl][student]
                pct_data[lbl][student]  = pct
                rank_data[lbl][student] = rank
            else:
                pct_data[lbl][student]  = np.nan   # 缺考
                rank_data[lbl][student] = np.nan

    pct_matrix  = pd.DataFrame(pct_data,  index=all_students)
    rank_matrix = pd.DataFrame(rank_data, index=all_students)

    # 合併班別班號：優先 scores.xlsx > analysis.xlsx 欄位
    merged_class = dict(class_from_excel)   # 先用 analysis.xlsx 的資料
    if class_info_df is not None:
        for _, row in class_info_df.iterrows():
            name = row['中文姓名']
            ban  = str(row.get('班別', '')).strip()
            num  = str(row.get('班號', '')).strip()
            if ban or num:
                merged_class[name] = (ban, num)  # scores.xlsx 優先覆蓋

    student_info = pd.DataFrame(
        [(s, *merged_class.get(s, ('', ''))) for s in all_students],
        columns=['中文姓名', '班別', '班號']
    )

    # 按班別（字串升序）→ 班號（數字升序）排序
    def _ban_num(x):
        try:    return int(float(str(x)))
        except: return 9999

    student_info = student_info.copy()
    student_info['_ban_sort'] = student_info['班別'].astype(str).str.strip()
    student_info['_num_sort'] = student_info['班號'].map(_ban_num)
    student_info = student_info.sort_values(['_ban_sort', '_num_sort']).drop(
        columns=['_ban_sort', '_num_sort'])
    student_info = student_info.reset_index(drop=True)

    # 同步矩陣行順序
    valid_students = [s for s in student_info['中文姓名'] if s in pct_matrix.index]
    pct_matrix  = pct_matrix.loc[valid_students]
    rank_matrix = rank_matrix.loc[valid_students]

    return pct_matrix, rank_matrix, exam_labels_sorted, student_info

# ══════════════════════════════════════════════════════════════
# 3. 趨勢計算
# ══════════════════════════════════════════════════════════════
def calc_trend(row_pct):
    """計算單個學生的趨勢指標"""
    valid = row_pct.dropna()
    if len(valid) < 2:
        return {'trend': '─', 'change': np.nan, 'avg': np.nan,
                'best': np.nan, 'worst': np.nan, 'attended': len(valid), 'absent': row_pct.isna().sum()}
    first, last = valid.iloc[0], valid.iloc[-1]
    change = last - first
    trend = '▲' if change > 2 else ('▼' if change < -2 else '─')
    return {
        'trend':   trend,
        'change':  round(change, 1),
        'avg':     round(valid.mean(), 1),
        'best':    round(valid.max(), 1),
        'worst':   round(valid.min(), 1),
        'attended': len(valid),
        'absent':  int(row_pct.isna().sum())
    }

def calc_class_stats(pct_matrix, pass_rate=0.4):
    """計算全班各考試統計"""
    stats = []
    for col in pct_matrix.columns:
        valid = pct_matrix[col].dropna()
        if len(valid) == 0:
            continue
        pass_count = (valid >= pass_rate * 100).sum()
        stats.append({
            '考試': col,
            '出席人數': len(valid),
            '缺考人數': int(pct_matrix[col].isna().sum()),
            '平均得分率(%)': round(valid.mean(), 1),
            '及格率(%)':    round(pass_count / len(valid) * 100, 1),
            '最高分(%)':    round(valid.max(), 1),
            '最低分(%)':    round(valid.min(), 1),
            '標準差':       round(valid.std(), 2),
        })
    return pd.DataFrame(stats)

# ══════════════════════════════════════════════════════════════
# 4. 圖表生成
# ══════════════════════════════════════════════════════════════
def make_class_trend_chart(class_stats_df, pass_rate=0.4, title="全班成績趨勢"):
    """生成全班趨勢圖（折線+柱狀）"""
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 7), dpi=120)
    fig.patch.set_facecolor('white')

    exams  = class_stats_df['考試'].tolist()
    avg    = class_stats_df['平均得分率(%)'].tolist()
    passrt = class_stats_df['及格率(%)'].tolist()
    x = range(len(exams))

    # 上圖：平均得分率
    bars = ax1.bar(x, avg, color='#1F6FB8', alpha=0.75, width=0.5, zorder=2)
    ax1.plot(x, avg, 'o-', color='#1A3A6B', linewidth=2, markersize=7, zorder=3)
    ax1.axhline(y=pass_rate * 100, color='orange', linestyle='--', linewidth=1.5, label=f'及格線 {int(pass_rate*100)}%')
    ax1.axhline(y=80, color='green', linestyle=':', linewidth=1.5, label='優秀線 80%')
    for i, v in enumerate(avg):
        ax1.text(i, v + 1.5, f'{v:.1f}%', ha='center', fontsize=10, fontweight='bold', color='#1A3A6B')
    ax1.set_ylim(0, 105)
    ax1.set_xticks(x)
    ax1.set_xticklabels(exams, fontsize=11)
    ax1.set_ylabel('平均得分率 (%)', fontsize=11)
    ax1.set_title(f'{title} — 平均得分率', fontsize=13, fontweight='bold', color='#1A3A6B')
    ax1.legend(fontsize=9)
    ax1.grid(axis='y', alpha=0.3)

    # 下圖：及格率
    colors = ['#177A3C' if v >= 60 else ('#E86B00' if v >= 40 else '#C0392B') for v in passrt]
    ax2.bar(x, passrt, color=colors, alpha=0.8, width=0.5, zorder=2)
    ax2.axhline(y=pass_rate * 100, color='orange', linestyle='--', linewidth=1.5)
    for i, v in enumerate(passrt):
        ax2.text(i, v + 1.5, f'{v:.1f}%', ha='center', fontsize=10, fontweight='bold')
    ax2.set_ylim(0, 105)
    ax2.set_xticks(x)
    ax2.set_xticklabels(exams, fontsize=11)
    ax2.set_ylabel('及格率 (%)', fontsize=11)
    ax2.set_title('全班及格率趨勢', fontsize=13, fontweight='bold', color='#1A3A6B')
    ax2.grid(axis='y', alpha=0.3)

    plt.tight_layout(pad=2.0)
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def make_student_trend_chart(student_name, pct_row, rank_row, pass_rate=0.4):
    """生成單個學生的成績趨勢圖"""
    exams  = pct_row.index.tolist()
    pcts   = pct_row.values
    ranks  = rank_row.values

    valid_mask = ~pd.isna(pcts)
    valid_x    = [i for i, v in enumerate(valid_mask) if v]
    valid_pcts = [pcts[i] for i in valid_x]

    if len(valid_pcts) == 0:
        return None

    fig, ax = plt.subplots(figsize=(9, 3.5), dpi=120)
    fig.patch.set_facecolor('white')

    # 柱狀背景
    bar_colors = []
    for p in pcts:
        if pd.isna(p):
            bar_colors.append('#E0E0E0')
        elif p >= 80:
            bar_colors.append('#D4EDDA')
        elif p >= pass_rate * 100:
            bar_colors.append('#D6E8F8')
        else:
            bar_colors.append('#FDECEA')

    ax.bar(range(len(exams)), [p if not pd.isna(p) else 0 for p in pcts],
           color=bar_colors, alpha=0.6, width=0.6, zorder=1)

    # 折線（只連有效點）
    if len(valid_x) >= 2:
        ax.plot(valid_x, valid_pcts, 'o-', color='#1F6FB8', linewidth=2.5,
                markersize=9, zorder=3)
    elif len(valid_x) == 1:
        ax.plot(valid_x, valid_pcts, 'o', color='#1F6FB8', markersize=10, zorder=3)

    # 分數標籤
    for i, (p, r) in enumerate(zip(pcts, ranks)):
        if pd.isna(p):
            ax.text(i, 3, '缺考', ha='center', fontsize=10, color='#888888', fontweight='bold')
        else:
            ax.text(i, p + 2, f'{p:.1f}%', ha='center', fontsize=10,
                    color='#1A3A6B', fontweight='bold')
            if not pd.isna(r):
                ax.text(i, p - 7, f'#{int(r)}', ha='center', fontsize=8.5,
                        color='#555555')

    ax.axhline(y=pass_rate * 100, color='orange', linestyle='--',
               linewidth=1.5, label=f'及格線 {int(pass_rate*100)}%', zorder=2)
    ax.axhline(y=80, color='green', linestyle=':',
               linewidth=1.5, label='優秀線 80%', zorder=2)

    ax.set_ylim(0, 108)
    ax.set_xticks(range(len(exams)))
    ax.set_xticklabels(exams, fontsize=11)
    ax.set_ylabel('得分率 (%)', fontsize=10)
    ax.set_title(f'{student_name} — 成績趨勢', fontsize=12, fontweight='bold', color='#1A3A6B')
    ax.legend(fontsize=8.5, loc='upper left')
    ax.grid(axis='y', alpha=0.3)

    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

# ══════════════════════════════════════════════════════════════
# 5. Excel 報告
# ══════════════════════════════════════════════════════════════
def apply_pct_color(ws, row, col, value, pass_rate=0.4):
    """根據得分率上色"""
    from openpyxl.styles import PatternFill, Font
    if value == '缺考' or (isinstance(value, float) and np.isnan(value)):
        ws.cell(row=row, column=col).value = '缺考'
        ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor="E0E0E0")
        ws.cell(row=row, column=col).font = Font(color="888888", italic=True)
    else:
        try:
            v = float(value)
            ws.cell(row=row, column=col).value = round(v, 1)
            if v >= 80:
                color = "D4EDDA"
            elif v >= pass_rate * 100:
                color = "D6E8F8"
            else:
                color = "FDECEA"
            ws.cell(row=row, column=col).fill = PatternFill("solid", fgColor=color)
        except:
            ws.cell(row=row, column=col).value = value

def export_tracking_excel(pct_matrix, rank_matrix, student_info, class_stats,
                           exam_labels, output_path, pass_rate=0.4, subject=""):
    """匯出成績追蹤 Excel 報告（4個工作表）"""
    from openpyxl import Workbook
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                  GradientFill)
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    navy  = "1A3A6B"
    blue  = "1F6FB8"
    white = "FFFFFF"
    lgray = "F2F2F2"

    header_font   = Font(bold=True, color=white, name="Microsoft JhengHei", size=11)
    normal_font   = Font(name="Microsoft JhengHei", size=10)
    bold_font     = Font(bold=True, name="Microsoft JhengHei", size=10)
    header_fill   = PatternFill("solid", fgColor=navy)
    subhdr_fill   = PatternFill("solid", fgColor=blue)
    alt_fill      = PatternFill("solid", fgColor=lgray)
    center_align  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align    = Alignment(horizontal="left",   vertical="center")
    thin_border   = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin")
    )

    # ── 工作表1：成績追蹤總表 ──
    ws1 = wb.active
    ws1.title = "1_📊 成績追蹤總表"

    title_row = f"DSE 成績追蹤報告 | {subject} | 共 {len(exam_labels)} 次考試 | 共 {len(pct_matrix)} 位學生"
    ws1.merge_cells(f"A1:{get_column_letter(4 + len(exam_labels) + 2)}1")
    c = ws1.cell(row=1, column=1, value=title_row)
    c.font = Font(bold=True, color=white, name="Microsoft JhengHei", size=12)
    c.fill = PatternFill("solid", fgColor=navy)
    c.alignment = center_align
    ws1.row_dimensions[1].height = 22

    headers = ["班別", "班號", "中文姓名"] + exam_labels + ["平均(%)", "趨勢", "變化(%)"]
    for j, h in enumerate(headers, 1):
        cell = ws1.cell(row=2, column=j, value=h)
        cell.font = header_font
        cell.fill = subhdr_fill if j > 3 else PatternFill("solid", fgColor=navy)
        cell.alignment = center_align
        cell.border = thin_border
    ws1.row_dimensions[2].height = 18

    for i, student in enumerate(pct_matrix.index):
        row_n = i + 3
        info = student_info[student_info["中文姓名"] == student]
        ban  = info["班別"].values[0] if len(info) else ""
        num  = info["班號"].values[0] if len(info) else ""
        fill = PatternFill("solid", fgColor="FFFFFF") if i % 2 == 0 else alt_fill

        for j, val in enumerate([ban, num, student], 1):
            cell = ws1.cell(row=row_n, column=j, value=val)
            cell.font = bold_font if j == 3 else normal_font
            cell.fill = fill
            cell.alignment = center_align if j != 3 else left_align
            cell.border = thin_border

        for k, exam in enumerate(exam_labels):
            apply_pct_color(ws1, row_n, 4 + k, pct_matrix.loc[student, exam], pass_rate)
            ws1.cell(row=row_n, column=4+k).alignment = center_align
            ws1.cell(row=row_n, column=4+k).border = thin_border

        t = calc_trend(pct_matrix.loc[student])
        avg_cell = ws1.cell(row=row_n, column=4+len(exam_labels),
                            value=t["avg"] if not np.isnan(t["avg"]) else "-")
        avg_cell.alignment = center_align
        avg_cell.border = thin_border

        trend_cell = ws1.cell(row=row_n, column=5+len(exam_labels), value=t["trend"])
        trend_cell.font = Font(bold=True, color=("177A3C" if t["trend"]=="▲"
                                                  else ("C0392B" if t["trend"]=="▼" else "888888")),
                               name="Microsoft JhengHei", size=12)
        trend_cell.alignment = center_align
        trend_cell.border = thin_border

        chg = t["change"]
        chg_val = f"+{chg:.1f}%" if (not np.isnan(chg) and chg > 0) else (f"{chg:.1f}%" if not np.isnan(chg) else "-")
        chg_cell = ws1.cell(row=row_n, column=6+len(exam_labels), value=chg_val)
        chg_cell.font = Font(color=("177A3C" if not np.isnan(chg) and chg > 0
                                     else ("C0392B" if not np.isnan(chg) and chg < 0 else "888888")),
                             name="Microsoft JhengHei", size=10)
        chg_cell.alignment = center_align
        chg_cell.border = thin_border

    # 欄寬
    ws1.column_dimensions["A"].width = 7
    ws1.column_dimensions["B"].width = 6
    ws1.column_dimensions["C"].width = 10
    for k in range(len(exam_labels)):
        ws1.column_dimensions[get_column_letter(4+k)].width = 12
    ws1.column_dimensions[get_column_letter(4+len(exam_labels))].width = 9
    ws1.column_dimensions[get_column_letter(5+len(exam_labels))].width = 7
    ws1.column_dimensions[get_column_letter(6+len(exam_labels))].width = 9

    # ── 工作表2：進退步分析 ──
    ws2 = wb.create_sheet("2_📈 進退步分析")
    h2 = ["班別","班號","中文姓名","出席次數","缺考次數","首次得分率(%)","最新得分率(%)","變化(%)","最高(%)","最低(%)","趨勢","平均(%)"]
    for j, h in enumerate(h2, 1):
        cell = ws2.cell(row=1, column=j, value=h)
        cell.font = header_font
        cell.fill = PatternFill("solid", fgColor=navy)
        cell.alignment = center_align
        cell.border = thin_border

    for i, student in enumerate(pct_matrix.index):
        row_n = i + 2
        info = student_info[student_info["中文姓名"] == student]
        ban  = info["班別"].values[0] if len(info) else ""
        num  = info["班號"].values[0] if len(info) else ""
        t    = calc_trend(pct_matrix.loc[student])
        valid = pct_matrix.loc[student].dropna()
        first_pct = round(valid.iloc[0], 1) if len(valid) > 0 else "-"
        last_pct  = round(valid.iloc[-1], 1) if len(valid) > 0 else "-"
        fill = PatternFill("solid", fgColor="FFFFFF") if i % 2 == 0 else alt_fill

        row_vals = [ban, num, student, t["attended"], t["absent"],
                    first_pct, last_pct,
                    f"+{t['change']:.1f}%" if (not np.isnan(t["change"]) and t["change"]>0) else
                    (f"{t['change']:.1f}%" if not np.isnan(t["change"]) else "-"),
                    t["best"] if not np.isnan(t["best"]) else "-",
                    t["worst"] if not np.isnan(t["worst"]) else "-",
                    t["trend"],
                    t["avg"] if not np.isnan(t["avg"]) else "-"]
        for j, v in enumerate(row_vals, 1):
            cell = ws2.cell(row=row_n, column=j, value=v)
            cell.font = normal_font
            cell.fill = fill
            cell.alignment = center_align
            cell.border = thin_border
            if j == 8:  # 變化%
                if isinstance(v, str) and v.startswith("+"):
                    cell.font = Font(color="177A3C", name="Microsoft JhengHei", size=10, bold=True)
                elif isinstance(v, str) and v.startswith("-"):
                    cell.font = Font(color="C0392B", name="Microsoft JhengHei", size=10, bold=True)

    for k, w in enumerate([7,6,10,8,8,12,12,10,10,10,7,10], 1):
        ws2.column_dimensions[get_column_letter(k)].width = w

    # ── 工作表3：全班趨勢 ──
    ws3 = wb.create_sheet("3_🏫 全班趨勢")
    h3 = list(class_stats.columns)
    for j, h in enumerate(h3, 1):
        cell = ws3.cell(row=1, column=j, value=h)
        cell.font = header_font
        cell.fill = PatternFill("solid", fgColor=navy)
        cell.alignment = center_align
        cell.border = thin_border

    for i, row_data in class_stats.iterrows():
        fill = PatternFill("solid", fgColor="FFFFFF") if i % 2 == 0 else alt_fill
        for j, v in enumerate(row_data, 1):
            cell = ws3.cell(row=i+2, column=j, value=v)
            cell.font = normal_font
            cell.fill = fill
            cell.alignment = center_align
            cell.border = thin_border

    for k, w in enumerate([14, 8, 8, 12, 10, 10, 10, 10], 1):
        ws3.column_dimensions[get_column_letter(k)].width = w

    # ── 工作表4：排名追蹤 ──
    ws4 = wb.create_sheet("4_🏆 排名追蹤")
    h4 = ["班別", "班號", "中文姓名"] + exam_labels + ["排名變化"]
    for j, h in enumerate(h4, 1):
        cell = ws4.cell(row=1, column=j, value=h)
        cell.font = header_font
        cell.fill = PatternFill("solid", fgColor=navy)
        cell.alignment = center_align
        cell.border = thin_border

    for i, student in enumerate(rank_matrix.index):
        row_n = i + 2
        info = student_info[student_info["中文姓名"] == student]
        ban  = info["班別"].values[0] if len(info) else ""
        num  = info["班號"].values[0] if len(info) else ""
        fill = PatternFill("solid", fgColor="FFFFFF") if i % 2 == 0 else alt_fill

        for j, v in enumerate([ban, num, student], 1):
            ws4.cell(row=row_n, column=j, value=v).font = normal_font
            ws4.cell(row=row_n, column=j).fill = fill
            ws4.cell(row=row_n, column=j).alignment = center_align
            ws4.cell(row=row_n, column=j).border = thin_border

        valid_ranks = rank_matrix.loc[student].dropna()
        for k, exam in enumerate(exam_labels):
            rv = rank_matrix.loc[student, exam]
            rv_num = pd.to_numeric(rv, errors="coerce")  # '-' / NaN / 數字 均轉數字
            display = '缺考' if pd.isna(rv_num) else int(rv_num)
            cell = ws4.cell(row=row_n, column=4+k, value=display)
            cell.fill = fill if display == "缺考" else PatternFill("solid",
                        fgColor=("FFD700" if display == 1 else
                                 ("C0C0C0" if display == 2 else
                                  ("CD7F32" if display == 3 else "FFFFFF"))))
            cell.alignment = center_align
            cell.border = thin_border

        # 排名變化（valid_ranks 已 dropna，再確保全為數值）
        valid_ranks_num = pd.to_numeric(valid_ranks, errors="coerce").dropna()
        if len(valid_ranks_num) >= 2:
            rank_chg = int(valid_ranks_num.iloc[0]) - int(valid_ranks_num.iloc[-1])
            chg_disp = f"↑{rank_chg}" if rank_chg > 0 else (f"↓{abs(rank_chg)}" if rank_chg < 0 else "─")
            chg_color = "177A3C" if rank_chg > 0 else ("C0392B" if rank_chg < 0 else "888888")
        else:
            chg_disp, chg_color = "-", "888888"

        chg_cell = ws4.cell(row=row_n, column=4+len(exam_labels), value=chg_disp)
        chg_cell.font = Font(color=chg_color, bold=True, name="Microsoft JhengHei", size=11)
        chg_cell.alignment = center_align
        chg_cell.border = thin_border

    for k, w in enumerate([7, 6, 10] + [10]*len(exam_labels) + [9], 1):
        ws4.column_dimensions[get_column_letter(k)].width = w

    wb.save(output_path)
    print(f"\n✅ Excel 追蹤報告已儲存：{output_path}")
    return output_path

# ══════════════════════════════════════════════════════════════
# 6. Word/PDF 追蹤報告
# ══════════════════════════════════════════════════════════════
def _set_font(run, name="Microsoft JhengHei", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*[int(color[i:i+2],16) for i in (0,2,4)])
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def generate_tracking_report(pct_matrix, rank_matrix, student_info,
                               class_stats, exam_labels, output_dir,
                               file_prefix, subject="", pass_rate=0.4):
    """生成 Word 追蹤報告並轉 PDF"""
    from docx2pdf import convert

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin  = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ── 封面 ──
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DSE 成績追蹤報告")
    _set_font(r, size=28, bold=True, color="1A3A6B")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subject if subject else file_prefix)
    _set_font(r2, size=18, color="1F6FB8")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"涵蓋 {len(exam_labels)} 次考試 · {len(pct_matrix)} 位學生")
    _set_font(r3, size=13, color="606060")
    doc.add_page_break()

    # ── 全班趨勢圖 ──
    p_hd = doc.add_paragraph()
    r_hd = p_hd.add_run("全班成績趨勢")
    _set_font(r_hd, size=16, bold=True, color="1A3A6B")

    class_chart_buf = make_class_trend_chart(class_stats, pass_rate, subject)
    doc.add_picture(class_chart_buf, width=Inches(6.5))
    class_chart_buf.close()
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 全班統計表
    tbl = doc.add_table(rows=1, cols=len(class_stats.columns))
    tbl.style = "Table Grid"
    for j, col_name in enumerate(class_stats.columns):
        cell = tbl.cell(0, j)
        _shade_cell(cell, "1A3A6B")
        p_ = cell.paragraphs[0]
        p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p_.add_run(col_name), size=9, bold=True, color="FFFFFF")

    for _, row_data in class_stats.iterrows():
        row = tbl.add_row()
        for j, v in enumerate(row_data):
            p_ = row.cells[j].paragraphs[0]
            p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_font(p_.add_run(str(v)), size=9)
    doc.add_page_break()

    # ── 個人成績追蹤 ──
    p_hd2 = doc.add_paragraph()
    r_hd2 = p_hd2.add_run("個人成績追蹤")
    _set_font(r_hd2, size=16, bold=True, color="1A3A6B")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    students_list = pct_matrix.index.tolist()
    for idx, student in enumerate(students_list, 1):
        info = student_info[student_info["中文姓名"] == student]
        ban  = info["班別"].values[0] if len(info) else ""
        num  = info["班號"].values[0] if len(info) else ""

        # 學生標題
        p_name = doc.add_paragraph()
        r_num  = p_name.add_run(f"  {idx:02d}  ")
        _set_font(r_num, size=10, bold=True, color="FFFFFF")
        r_name = p_name.add_run(f"  {ban}{int(num):02d} {student}" if str(num).isdigit() else f"  {student}")
        _set_font(r_name, size=13, bold=True, color="1A3A6B")

        # 趨勢圖
        chart_buf = make_student_trend_chart(
            student, pct_matrix.loc[student], rank_matrix.loc[student], pass_rate)
        if chart_buf:
            doc.add_picture(chart_buf, width=Inches(6.2))
            chart_buf.close()

        # 數據摘要行
        t = calc_trend(pct_matrix.loc[student])
        valid = pct_matrix.loc[student].dropna()
        summary_parts = []
        if len(valid) > 0:
            summary_parts.append(f"出席：{t['attended']} 次")
        if t["absent"] > 0:
            summary_parts.append(f"缺考：{t['absent']} 次")
        if not np.isnan(t["avg"]):
            summary_parts.append(f"平均：{t['avg']}%")
        if not np.isnan(t["change"]):
            chg_str = f"+{t['change']}%" if t["change"] > 0 else f"{t['change']}%"
            summary_parts.append(f"變化：{chg_str}")

        p_sum = doc.add_paragraph("  " + "　|　".join(summary_parts))
        p_sum.paragraph_format.space_after = Pt(8)

        if idx % 3 == 0 and idx < len(students_list):
            doc.add_page_break()

    # 儲存 DOCX
    os.makedirs(output_dir, exist_ok=True)
    docx_path = os.path.join(output_dir, f"{file_prefix}_成績追蹤報告.docx")
    doc.save(docx_path)
    print(f"✅ Word 追蹤報告已儲存：{docx_path}")

    # 轉 PDF
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        convert(docx_path, pdf_path)
        print(f"✅ PDF 追蹤報告已儲存：{pdf_path}")
        return docx_path, pdf_path
    except Exception as e:
        print(f"⚠️  PDF 轉換失敗（Word 文件已儲存）：{e}")
        return docx_path, None

def generate_tracking_report_bytes(pct_matrix, rank_matrix, student_info,
                                    class_stats, exam_labels,
                                    file_prefix, subject="", pass_rate=0.4):
    """生成 Word 追蹤報告並回傳 (docx_bytes, pdf_bytes_or_None)"""
    import io as _io, subprocess, tempfile, os as _os

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7);  section.page_width  = Cm(21.0)
    section.top_margin  = Cm(1.8);   section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2);   section.right_margin = Cm(2.2)

    # ── 封面 ──
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("DSE 成績追蹤報告"), size=28, bold=True, color="1A3A6B")
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p2.add_run(subject if subject else file_prefix), size=18, color="1F6FB8")
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p3.add_run(f"涵蓋 {len(exam_labels)} 次考試 · {len(pct_matrix)} 位學生"),
              size=13, color="606060")
    doc.add_page_break()

    # ── 全班趨勢圖 ──
    p_hd = doc.add_paragraph()
    _set_font(p_hd.add_run("全班成績趨勢"), size=16, bold=True, color="1A3A6B")
    class_chart_buf = make_class_trend_chart(class_stats, pass_rate, subject)
    doc.add_picture(class_chart_buf, width=Inches(6.5)); class_chart_buf.close()
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 全班統計表
    tbl = doc.add_table(rows=1, cols=len(class_stats.columns)); tbl.style = "Table Grid"
    for j, col_name in enumerate(class_stats.columns):
        cell = tbl.cell(0, j); _shade_cell(cell, "1A3A6B")
        p_ = cell.paragraphs[0]; p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p_.add_run(col_name), size=9, bold=True, color="FFFFFF")
    for _, row_data in class_stats.iterrows():
        row = tbl.add_row()
        for j, v in enumerate(row_data):
            p_ = row.cells[j].paragraphs[0]; p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_font(p_.add_run(str(v)), size=9)
    doc.add_page_break()

    # ── 個人成績追蹤 ──
    p_hd2 = doc.add_paragraph()
    _set_font(p_hd2.add_run("個人成績追蹤"), size=16, bold=True, color="1A3A6B")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    students_list = pct_matrix.index.tolist()
    for idx, student in enumerate(students_list, 1):
        info = student_info[student_info["中文姓名"] == student]
        ban  = info["班別"].values[0] if len(info) else ""
        num  = info["班號"].values[0]  if len(info) else ""
        p_name = doc.add_paragraph()
        _set_font(p_name.add_run(f"  {idx:02d}  "), size=10, bold=True, color="FFFFFF")
        _set_font(p_name.add_run(
            f"  {ban}{int(num):02d} {student}" if str(num).isdigit() else f"  {student}"),
            size=13, bold=True, color="1A3A6B")
        chart_buf = make_student_trend_chart(
            student, pct_matrix.loc[student], rank_matrix.loc[student], pass_rate)
        if chart_buf:
            doc.add_picture(chart_buf, width=Inches(6.2)); chart_buf.close()
        t = calc_trend(pct_matrix.loc[student])
        valid = pct_matrix.loc[student].dropna()
        parts = []
        if len(valid) > 0:      parts.append(f"出席：{t['attended']} 次")
        if t["absent"] > 0:     parts.append(f"缺考：{t['absent']} 次")
        if not np.isnan(t["avg"]):    parts.append(f"平均：{t['avg']}%")
        if not np.isnan(t["change"]):
            chg = f"+{t['change']}%" if t["change"] > 0 else f"{t['change']}%"
            parts.append(f"變化：{chg}")
        p_sum = doc.add_paragraph("  " + "　|　".join(parts))
        p_sum.paragraph_format.space_after = Pt(8)
        if idx % 3 == 0 and idx < len(students_list):
            doc.add_page_break()

    # 存為 bytes
    docx_buf = _io.BytesIO(); doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    # 轉 PDF（LibreOffice）
    pdf_bytes = None
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = _os.path.join(tmpdir, "tracking.docx")
        pdf_path  = _os.path.join(tmpdir, "tracking.pdf")
        with open(docx_path, "wb") as f: f.write(docx_bytes)
        for lo in ["libreoffice", "/usr/bin/libreoffice",
                   "/usr/lib/libreoffice/program/soffice"]:
            try:
                subprocess.run([lo, "--headless", "--convert-to", "pdf",
                                "--outdir", tmpdir, docx_path],
                               capture_output=True, timeout=120)
                if _os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f: pdf_bytes = f.read()
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

    return docx_bytes, pdf_bytes


# ══════════════════════════════════════════════════════════════
# 7. 主程式
# ══════════════════════════════════════════════════════════════
def main():
    print("=" * 60)
    print("  📊 DSE 成績追蹤系統  v1.0")
    print("=" * 60)

    # 讀取 scores.xlsx 取得班別班號資訊
    class_info_df = None
    if os.path.exists("scores.xlsx"):
        try:
            raw = pd.read_excel("scores.xlsx", header=None)
            ci = pd.DataFrame()
            ci["班別"]    = raw.iloc[4:, 0].values
            ci["班號"]    = raw.iloc[4:, 1].values
            ci["中文姓名"] = raw.iloc[4:, 3].values
            ci = ci[ci["中文姓名"].notna() &
                    ~ci["中文姓名"].astype(str).str.contains("說明|輸入", na=False)]
            class_info_df = ci
            print(f"✅ 讀取 scores.xlsx：{len(ci)} 位學生的班別班號資訊")
        except Exception as e:
            print(f"⚠️  讀取 scores.xlsx 失敗：{e}")

    # 掃描考試檔案（先全部掃描，再按年級/科目過濾）
    print("\n🔍 掃描 *_analysis.xlsx 檔案...")
    all_exam_files = scan_analysis_files(".")
    if not all_exam_files:
        print("❌ 未找到任何 *_analysis.xlsx 檔案！")
        print("   請先執行 exam_item_analysis.py 生成至少一份分析報告")
        input("\n按 Enter 退出...")
        return

    print(f"\n   找到 {len(all_exam_files)} 個分析檔案：")
    forms_found    = sorted(set(ef['form']    for ef in all_exam_files if ef['form']))
    subjects_found = sorted(set(ef['subject'] for ef in all_exam_files if ef['subject']))
    for ef in all_exam_files:
        tag = " ".join(filter(None, [ef['form'], ef['subject']]))
        print(f"   · {os.path.basename(ef['file'])}  [{tag or '舊格式'}]")

    # 年級過濾（偵測到多個年級時才詢問）
    form_filter = None
    if len(forms_found) > 1:
        print(f"\n   偵測到多個年級：{', '.join(forms_found)}")
        fi = input("   請輸入年級過濾（如 F5），按 Enter 不過濾：").strip().upper()
        form_filter = fi or None
    elif len(forms_found) == 1:
        form_filter = forms_found[0]
        print(f"   ✅ 只有一個年級 {form_filter}，自動套用")

    # 科目過濾（偵測到多個科目時才詢問）
    subject_filter = None
    if len(subjects_found) > 1:
        print(f"\n   偵測到多個科目：{', '.join(subjects_found)}")
        si = input("   請輸入科目過濾（如 BAFS），按 Enter 不過濾：").strip()
        subject_filter = si or None
    elif len(subjects_found) == 1:
        subject_filter = subjects_found[0]
        print(f"   ✅ 只有一個科目 {subject_filter}，自動套用")

    exam_files = scan_analysis_files(".", filter_form=form_filter, filter_subject=subject_filter)
    if not exam_files:
        print("❌ 過濾後找不到符合條件的檔案，請確認年級/科目名稱是否正確")
        input("\n按 Enter 退出...")
        return
    print(f"\n   ✅ 過濾後共 {len(exam_files)} 個檔案納入分析")
    for ef in exam_files:
        print(f"   · {os.path.basename(ef['file'])}  ({ef['year']} {ef['type']})")

    # 科目名稱（用於報告標題，若已從過濾得到則直接使用）
    subject = subject_filter or input("\n📚 請輸入科目名稱（選填，按 Enter 跳過）：").strip()

    # 及格線
    print("\n📊 請選擇及格分數線：")
    print("   [1] 40%（適用：高中科目）")
    print("   [2] 50%（適用：初中科目）")
    while True:
        pi = input("   請輸入選項（1 或 2，預設 1）：").strip()
        if pi in ["", "1"]:
            pass_rate = 0.4; break
        elif pi == "2":
            pass_rate = 0.5; break
        else:
            print("   ⚠️  請輸入 1 或 2")

    # 輸出檔名前綴
    years = sorted(set(ef["year"] for ef in exam_files))
    year_str = "_".join(years)
    types = "_".join(ef["type"] for ef in exam_files)
    file_prefix = f"{year_str}_{types}_追蹤" if len(exam_files) <= 4 else f"{year_str}_全追蹤"

    # 建立矩陣
    print("\n📐 建立成績追蹤矩陣...")
    pct_matrix, rank_matrix, exam_labels, student_info = build_tracking_matrix(
        exam_files, class_info_df)

    if pct_matrix is None:
        print("❌ 無法建立成績矩陣，請檢查檔案格式")
        input("\n按 Enter 退出...")
        return

    print(f"\n   ✅ 矩陣大小：{pct_matrix.shape[0]} 位學生 × {pct_matrix.shape[1]} 次考試")
    absent_count = pct_matrix.isna().sum().sum()
    print(f"   缺考記錄：{int(absent_count)} 筆")

    # 全班統計
    class_stats = calc_class_stats(pct_matrix, pass_rate)

    print("\n全班各考試統計：")
    print(class_stats.to_string(index=False))

    # 匯出 Excel
    print("\n💾 匯出 Excel 追蹤報告...")
    excel_path = f"{file_prefix}_成績追蹤.xlsx"
    export_tracking_excel(pct_matrix, rank_matrix, student_info, class_stats,
                           exam_labels, excel_path, pass_rate, subject)

    # 生成 PDF 報告
    gen = input("\n📄 是否生成 PDF 追蹤報告（含個人趨勢圖）？(Y/N)：").strip().upper()
    if gen == "Y":
        print("\n生成中（請稍候）...")
        docx_p, pdf_p = generate_tracking_report(
            pct_matrix, rank_matrix, student_info, class_stats,
            exam_labels, ".", file_prefix, subject, pass_rate)

    print(f"""
╔══════════════════════════════════════════╗
║  ✅ 成績追蹤完成！                       ║
╠══════════════════════════════════════════╣
║  📊 Excel：{excel_path:<31}║
║  📄 PDF：{(file_prefix+"_成績追蹤報告.pdf"):<33}║
╚══════════════════════════════════════════╝
""")
    input("按 Enter 退出...")

if __name__ == "__main__":
    main()
