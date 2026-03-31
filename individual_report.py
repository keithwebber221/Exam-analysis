# individual_report.py
# 個人報告生成器 v2.4（增強版）- 大圖表、雷達圖、美觀排版
# 依賴：pip install python-docx pillow matplotlib numpy pandas scipy

import os
import sys
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import matplotlib
matplotlib.use('Agg')
from matplotlib import pyplot as plt
from matplotlib import rcParams
from matplotlib import font_manager
import sys

def _setup_cjk_font():
    """動態偵測並設定中文字型（支援 macOS / Windows / Linux Streamlit Cloud）"""
    candidates = [
        # Linux / Streamlit Cloud（需 packages.txt: fonts-noto-cjk）
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJKtc-Regular.otf",
        # macOS
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        # Windows
        "C:/Windows/Fonts/msjh.ttc",
        "C:/Windows/Fonts/mingliu.ttc",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                font_manager.fontManager.addfont(path)
                prop = font_manager.FontProperties(fname=path)
                name = prop.get_name()
                rcParams['font.family']     = 'sans-serif'
                rcParams['font.sans-serif'] = [name] + rcParams.get('font.sans-serif', [])
                rcParams['axes.unicode_minus'] = False
                return name
            except Exception:
                continue
    rcParams['font.family']        = 'sans-serif'
    rcParams['axes.unicode_minus'] = False
    return 'DejaVu Sans'

DEFAULT_FONT = _setup_cjk_font()
rcParams['axes.unicode_minus'] = False


def set_cell_background(cell, color):
    """設定單元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_large_score_chart(df_student, max_scores):
    """生成放大的各題成績對比棒圖"""
    try:
        fig, ax = plt.subplots(figsize=(11, 3.5), dpi=100)

        questions = list(df_student.index)
        scores = [df_student[q] for q in questions]
        max_s = [max_scores[q] for q in questions]
        rates = [scores[i]/max_s[i]*100 if max_s[i] > 0 else 0 for i in range(len(questions))]

        x = np.arange(len(questions))
        width = 0.35

        bars1 = ax.bar(x - width/2, scores, width, label='得分', color='#3498db', alpha=0.85, edgecolor='navy', linewidth=1)
        bars2 = ax.bar(x + width/2, max_s, width, label='滿分', color='#ecf0f1', alpha=0.85, edgecolor='gray', linewidth=1)

        # 在柱頂加上數值和得分率
        for i, (bar, score, max_sc, rate) in enumerate(zip(bars1, scores, max_s, rates)):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.3,
                   f'{int(score)}\n{rate:.0f}%', ha='center', va='bottom', fontsize=10, fontweight='bold')

        ax.set_ylabel('分數', fontsize=12, fontweight='bold')
        ax.set_title('各題成績對比分析', fontsize=14, fontweight='bold', pad=15)
        ax.set_xticks(x)
        ax.set_xticklabels(questions, fontsize=11, fontweight='bold')
        ax.legend(fontsize=11, loc='upper right')
        ax.grid(axis='y', alpha=0.3, linestyle='--', linewidth=0.8)
        ax.set_axisbelow(True)

        fig.patch.set_facecolor('white')
        plt.tight_layout()

        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return buf
    except Exception as e:
        print(f"   ⚠️  各題成績圖表生成失敗：{e}")
        return None


def create_single_bar_progress(student_total, total_max, class_avg_scores, class_total_max):
    """生成單條棒進度圖表（個人 vs 全班平均）"""
    try:
        fig, ax = plt.subplots(figsize=(11, 2), dpi=100)

        student_rate = student_total / total_max * 100 if total_max > 0 else 0
        class_avg_rate = class_avg_scores / class_total_max * 100 if class_total_max > 0 else 0

        # 背景棒（全分）
        ax.barh(1, 100, height=0.3, color='#ecf0f1', alpha=0.5, edgecolor='gray', linewidth=1)

        # 全班平均棒
        ax.barh(1, class_avg_rate, height=0.25, color='#f39c12', alpha=0.7, edgecolor='#d68910', linewidth=2, label=f'全班平均 ({class_avg_rate:.1f}%)')

        # 個人得分棒
        ax.barh(1, student_rate, height=0.15, color='#3498db', alpha=0.95, edgecolor='navy', linewidth=2, label=f'個人成績 ({student_rate:.1f}%)')

        # 標籤
        ax.text(student_rate + 2, 1.08, f'{int(student_total)}/{int(total_max)}', 
               fontsize=11, fontweight='bold', color='#3498db', va='center')
        ax.text(class_avg_rate + 2, 0.92, f'{int(class_avg_scores)}/{int(class_total_max)}', 
               fontsize=10, fontweight='bold', color='#f39c12', va='center')

        # 參考線
        ax.axvline(x=80, color='green', linestyle='--', alpha=0.4, linewidth=2)
        ax.text(80, 1.35, '80% 優秀線', fontsize=10, color='green', fontweight='bold', ha='center')

        ax.set_xlim(0, 110)
        ax.set_ylim(0.5, 1.5)
        ax.set_xlabel('得分率 %', fontsize=12, fontweight='bold')
        ax.set_title('個人成績與全班平均比較', fontsize=14, fontweight='bold', pad=15)
        ax.set_yticks([])
        ax.legend(fontsize=11, loc='lower right')
        ax.grid(axis='x', alpha=0.3, linestyle='--', linewidth=0.8)
        ax.set_axisbelow(True)

        fig.patch.set_facecolor('white')
        plt.tight_layout()

        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        plt.close(fig)
        return buf
    except Exception as e:
        print(f"   ⚠️  進度條圖表生成失敗：{e}")
        return None



def create_strong_weak_analysis_chart_v2(df_student, max_scores, pass_rate=0.4):
    """生成改進的強弱項分析圖表"""
    try:
        # 計算強弱項
        strong_items = []
        good_items = []
        weak_items = []

        for q in df_student.index:
            score = df_student[q]
            max_s = max_scores[q]
            rate = score / max_s * 100 if max_s > 0 else 0

            if rate >= 80:
                strong_items.append((q, rate))
            elif rate >= pass_rate * 100:
                good_items.append((q, rate))
            else:
                weak_items.append((q, rate))

        # 準備數據
        if strong_items or good_items or weak_items:
            all_items = strong_items + good_items + weak_items
            labels = [item[0] for item in all_items]
            rates = [item[1] for item in all_items]

            # 顏色映射
            colors = []
            for r in rates:
                if r >= 80:
                    colors.append('#2ecc71')  # 綠色 - 強項
                elif r >= pass_rate * 100:
                    colors.append('#f39c12')  # 橙色 - 良好
                else:
                    colors.append('#e74c3c')  # 紅色 - 弱項

            fig, ax = plt.subplots(figsize=(11, 3.5), dpi=100)

            bars = ax.bar(range(len(labels)), rates, color=colors, alpha=0.8, width=0.6, edgecolor='black', linewidth=1)

            # 在柱頂加上百分比和數值
            for bar, rate in zip(bars, rates):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 2,
                       f'{rate:.0f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')

            ax.set_ylabel('得分率 %', fontsize=12, fontweight='bold')
            ax.set_title('各題得分率分級分析', fontsize=14, fontweight='bold', pad=15)
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontsize=11, fontweight='bold')
            ax.set_ylim(0, 110)
            ax.grid(axis='y', alpha=0.3, linestyle='--', linewidth=0.8)
            ax.set_axisbelow(True)

            # 加上分級線
            ax.axhline(y=80, color='green', linestyle='--', alpha=0.4, linewidth=1.5)
            ax.axhline(y=pass_rate * 100, color='orange', linestyle='--', alpha=0.4, linewidth=1.5)
            ax.text(-1, 82, '優秀 ≥80%', fontsize=9, color='green', fontweight='bold')
            ax.text(-1, pass_rate * 100 + 2, f'及格 ≥{int(pass_rate*100)}%', fontsize=9, color='orange', fontweight='bold')

            fig.patch.set_facecolor('white')
            plt.tight_layout()

            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=100, bbox_inches='tight', facecolor='white')
            buf.seek(0)
            plt.close(fig)
            return buf, strong_items, good_items, weak_items
        else:
            plt.close('all')
            return None, [], [], []
    except Exception as e:
        print(f"   ⚠️  強弱項圖表生成失敗：{e}")
        return None, [], [], []


def docx_to_pdf(docx_path, pdf_path):
    """DOCX 轉 PDF：優先使用 LibreOffice headless（macOS/Linux/Windows 均無需授權彈窗）"""
    import subprocess, sys, shutil

    pdf_dir  = os.path.dirname(os.path.abspath(pdf_path))
    pdf_dir  = pdf_dir if pdf_dir else "."

    # LibreOffice 執行檔候選路徑（macOS App、Homebrew、Linux、Windows）
    lo_candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS App
        shutil.which("soffice") or "",                            # PATH (brew/linux)
        shutil.which("libreoffice") or "",                        # PATH 別名
        r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
    ]
    lo_bin = next((p for p in lo_candidates if p and os.path.exists(p)), None)

    if lo_bin:
        try:
            # LibreOffice 輸出檔名 = 原 docx 同名改副檔名 .pdf
            expected_pdf = os.path.join(
                pdf_dir,
                os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            )
            result = subprocess.run(
                [lo_bin, "--headless", "--convert-to", "pdf",
                 "--outdir", pdf_dir, os.path.abspath(docx_path)],
                capture_output=True, timeout=60
            )
            # 若 LibreOffice 輸出路徑與預期 pdf_path 不同，重新命名
            if os.path.exists(expected_pdf) and os.path.abspath(expected_pdf) != os.path.abspath(pdf_path):
                os.replace(expected_pdf, pdf_path)
            if os.path.exists(pdf_path):
                return True
        except Exception:
            pass

    # Fallback：docx2pdf（僅 Windows/Linux 使用，macOS 會觸發授權彈窗）
    if sys.platform != "darwin":
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return True
        except Exception:
            pass

    return False


def create_personal_report_v2_4(student_name, total_score, total_max,
                               df_student_scores, max_scores, item_df, exam_info, 
                               class_avg_scores, class_total_max, class_info=None,
                               pass_rate=0.4):
    """生成新版個人報告 v2.5（精簡版）"""

    doc = Document()

    # 設定 A4 直向
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    # ========== 標題區域 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(exam_info["exam_title"])
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = DEFAULT_FONT
    title_run.font.color.rgb = RGBColor(25, 25, 112)

    # 副標題
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("學生個人成績報告")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.bold = True
    subtitle_run.font.name = DEFAULT_FONT

    # 學生基本資訊（放大）- 包含班別和班號
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 從 class_info 讀取班別/班號
    class_name = 'N/A'
    class_num = 'N/A'
    if class_info is not None:
        match = class_info[class_info["中文姓名"] == student_name]
        if not match.empty:
            class_name = str(match.iloc[0]["班別"])
            class_num = str(match.iloc[0]["班號"])

    class_run = info_para.add_run(f"班別：{class_name}  班號：{class_num}  ")
    class_run.font.size = Pt(11)
    class_run.font.bold = True
    class_run.font.name = DEFAULT_FONT

    student_name_run = info_para.add_run(f"姓名：{student_name}  ")
    student_name_run.font.size = Pt(11)
    student_name_run.font.bold = True
    student_name_run.font.name = DEFAULT_FONT

    score_run = info_para.add_run(f"總分：{int(total_score)}/{int(total_max)}  ({total_score/total_max*100:.1f}%)")
    score_run.font.size = Pt(11)
    score_run.font.bold = True
    score_run.font.color.rgb = RGBColor(0, 0, 0)
    score_run.font.name = DEFAULT_FONT

    doc.add_paragraph()  # 空行

    # ========== 成績表格（簡化，只保留必要資訊） ==========
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    # 設定列寬
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(1.0)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "題號"
    hdr_cells[1].text = "滿分"
    hdr_cells[2].text = "得分"
    hdr_cells[3].text = "得分率"

    for cell in hdr_cells:
        set_cell_background(cell, "1f4e78")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = DEFAULT_FONT
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for q in df_student_scores.index:
        score = df_student_scores[q]
        max_s = max_scores[q]
        rate = score / max_s * 100 if max_s > 0 else 0

        row_cells = table.add_row().cells
        row_cells[0].text = str(q)
        row_cells[1].text = str(int(max_s))
        row_cells[2].text = str(int(score))
        row_cells[3].text = f"{rate:.1f}%"

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = DEFAULT_FONT
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # ========== 圖表區域 ==========

    # 進度條（個人 vs 全班平均 - 單條棒）
    try:
        progress_buf = create_single_bar_progress(total_score, total_max, class_avg_scores, class_total_max)
        if progress_buf:
            progress_para = doc.add_paragraph()
            progress_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(progress_buf, width=Inches(7.5))
            progress_buf.close()
    except:
        pass

    doc.add_paragraph()  # 空行

    # 圖 3：各題得分率分級分析
    analysis_buf, strong_items, good_items, weak_items = create_strong_weak_analysis_chart_v2(df_student_scores, max_scores, pass_rate)
    if analysis_buf:
        analysis_para = doc.add_paragraph()
        analysis_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(analysis_buf, width=Inches(7.5))
        analysis_buf.close()

    return doc






def merge_individual_pdfs(pdf_folder, class_info, file_prefix):
    """
    將所有個人報告 PDF 合併成單一檔案

    Parameters:
    -----------
    pdf_folder : str
        個人報告 PDF 存放目錄
    class_info : DataFrame
        班級資訊（含班號用於排序）
    file_prefix : str
        檔案前綴（用於輸出檔名）

    Returns:
    --------
    merged_pdf_path : str
        合併後 PDF 的路徑；若失敗返回 None
    """

    try:
        from PyPDF2 import PdfMerger
    except ImportError:
        try:
            from pypdf import PdfMerger
        except ImportError:
            print("⚠️  PDF 合併需要 PyPDF2 或 pypdf 庫")
            print("   執行：pip install PyPDF2")
            return None

    import glob

    # 獲取所有 PDF 檔案
    pdf_files = glob.glob(os.path.join(pdf_folder, "*_個人報告.pdf"))

    if not pdf_files:
        print(f"⚠️  在 {pdf_folder} 中未找到個人報告 PDF")
        return None

    # 建立學生名稱到班別、班號的映射
    name_to_class_info = {}
    for _, row in class_info.iterrows():
        name = row['中文姓名']
        class_code = row.get('班別', 'ZZ')
        try:
            class_num = int(row['班號'])
        except (ValueError, TypeError):
            class_num = 999  # 無效班號排最後
        name_to_class_info[name] = (class_code, class_num)

    # 排序 PDF 檔案（按班別、班號）
    def get_sort_key(pdf_path):
        # 從檔名提取班別、班號、學生名稱
        # 檔名格式：5A01李栢彤_個人報告.pdf 或 李栢彤_個人報告.pdf (舊格式)
        filename = os.path.basename(pdf_path)
        student_part = filename.replace('_個人報告.pdf', '')

        # 嘗試從新格式提取班別班號（前5個字符通常是班別班號）
        class_code_from_name = ""
        class_num_from_name = 0

        # 檢查是否是新格式（以班別開頭，如 5A01）
        if len(student_part) >= 4:
            potential_class = student_part[:4]
            # 簡單檢查：第一個是數字，第二個是字母，第3-4個是數字
            try:
                if (potential_class[0].isdigit() and 
                    potential_class[1].isalpha() and 
                    potential_class[2:].isdigit()):
                    class_code_from_name = potential_class[:2]  # 如 "5A"
                    class_num_from_name = int(potential_class[2:])  # 如 01
                    student_name = student_part[4:]  # 剩下的是名字
                else:
                    student_name = student_part
            except:
                student_name = student_part
        else:
            student_name = student_part

        # 若從檔名成功提取，使用提取的值；否則查表
        if class_code_from_name and class_num_from_name > 0:
            return (class_code_from_name, class_num_from_name, student_name)
        elif student_name in name_to_class_info:
            class_code, class_num = name_to_class_info[student_name]
            return (class_code, class_num, student_name)
        else:
            return ("ZZ", 999, student_name)

    pdf_files_sorted = sorted(pdf_files, key=get_sort_key)

    print(f"▶ 合併 {len(pdf_files_sorted)} 份個人報告 PDF...")
    for i, pdf_path in enumerate(pdf_files_sorted, 1):
        filename = os.path.basename(pdf_path)
        print(f"  {i}. {filename}")

    # 合併 PDF
    try:
        merger = PdfMerger()

        for pdf_path in pdf_files_sorted:
            try:
                merger.append(pdf_path)
            except Exception as e:
                print(f"  ⚠️  無法添加 {os.path.basename(pdf_path)}：{e}")
                continue

        # 輸出合併檔案
        merged_pdf_path = os.path.join(
            os.path.dirname(pdf_folder),
            f"{file_prefix}_個人報告合併.pdf"
        )

        merger.write(merged_pdf_path)
        merger.close()

        merged_size = os.path.getsize(merged_pdf_path)
        print(f"\n✅ PDF 合併完成！")
        print(f"   檔案：{merged_pdf_path}")
        print(f"   大小：{merged_size/1024:.1f} KB")

        return merged_pdf_path

    except Exception as e:
        print(f"❌ PDF 合併失敗：{e}")
        import traceback
        traceback.print_exc()
        return None



def generate_combined_class_report(df, max_scores, item_df, exam_info, class_info, output_dir, pass_rate=0.4, absent_set=None):
    """生成全班整合報告（所有學生報告合成一份）"""
    absent_set = absent_set or set()

    os.makedirs(output_dir, exist_ok=True)

    # 保存 class_info 供 PDF 合併使用
    df_class_info = class_info.copy()

    # 建立主文件
    combined_doc = Document()

    # ========== 設定頁面邊距 ==========
    sections = combined_doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ========== 封面 ==========
    title = combined_doc.add_heading(exam_info.get('exam_title', '考試成績報告'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(25, 25, 112)

    subtitle = combined_doc.add_paragraph('全班成績報告')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True

    # 報告日期
    date_str = datetime.now().strftime("%Y年%m月%d日")
    date_para = combined_doc.add_paragraph(f'報告生成日期：{date_str}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 班級資訊統計
    class_info_para = combined_doc.add_paragraph()
    run = class_info_para.add_run('班級資訊')
    run.bold = True
    class_info_para.add_run(f'\n學生人數：{len(df)} 人\n試卷滿分：{int(max_scores.sum())} 分\n試題數量：{len(max_scores)} 題')

    combined_doc.add_paragraph()

    # ========== 班級統計摘要 ==========
    combined_doc.add_heading('班級成績統計', 1)

    total_scores = df.sum(axis=1)
    total_max = int(max_scores.sum())

    stats_para = combined_doc.add_paragraph()
    run = stats_para.add_run('成績統計')
    run.bold = True
    stats_para.add_run(f'\n平均分：{total_scores.mean():.1f}/{total_max}')
    stats_para.add_run(f'\n最高分：{total_scores.max():.1f}/{total_max}')
    stats_para.add_run(f'\n最低分：{total_scores.min():.1f}/{total_max}')
    stats_para.add_run(f'\n標準差：{total_scores.std():.2f}')

    # 成績分佈
    passing_count = len([s for s in total_scores if s >= total_max * pass_rate])
    excellent_count = len([s for s in total_scores if s >= total_max * 0.8])

    run = stats_para.add_run('\n成績達標分析')
    run.bold = True
    stats_para.add_run(f'\n優秀（≥80%）：{excellent_count} 人')
    stats_para.add_run(f'\n及格（≥{int(pass_rate*100)}%）：{passing_count} 人')

    combined_doc.add_paragraph()

    # ========== 各題分析 ==========
    combined_doc.add_heading('試題難度分析', 1)

    # 建立試題表格
    table = combined_doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    # 表頭
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '題號'
    hdr_cells[1].text = '滿分'
    hdr_cells[2].text = '平均分'
    hdr_cells[3].text = '難度'

    # 表頭格式
    for cell in hdr_cells:
        set_cell_background(cell, '1F2150')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    # 填充試題資料
    for question in item_df.index:
        row_cells = table.add_row().cells
        max_s = max_scores[question]
        avg_s = item_df.loc[question, '平均分']
        difficulty = 100 - (avg_s / max_s * 100) if max_s > 0 else 0

        row_cells[0].text = str(question)
        row_cells[1].text = str(int(max_s))
        row_cells[2].text = f"{avg_s:.1f}"

        if difficulty >= 70:
            row_cells[3].text = '難'
        elif difficulty >= 40:
            row_cells[3].text = '中'
        else:
            row_cells[3].text = '易'

    combined_doc.add_paragraph()
    combined_doc.add_page_break()

    # ========== 個人報告部分 ==========
    combined_doc.add_heading('個別學生成績報告', 1)
    combined_doc.add_paragraph()

    # 按班別、班號排序學生
    total_scores_dict = total_scores.to_dict()

    def get_student_sort_key(name):
        row = class_info[class_info['中文姓名'] == name]
        if not row.empty:
            class_code = row.iloc[0].get('班別', 'ZZ')
            try:
                class_num = int(row.iloc[0]['班號'])
            except (ValueError, TypeError):
                class_num = 999
            return (class_code, class_num)
        return ('ZZ', 999)

    sorted_students = sorted(
        class_info['中文姓名'].tolist(),
        key=get_student_sort_key
    )

    # 逐個新增學生報告
    for idx, student_name in enumerate(sorted_students):
        if student_name not in df.index:
            continue

        # 學生標題
        student_heading = combined_doc.add_heading(str(student_name), 2)
        student_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 班級資訊
        student_row = class_info[class_info['中文姓名'] == student_name].iloc[0]
        info_para = combined_doc.add_paragraph()
        info_para.add_run(f"班別：{student_row['班別']} | 班號：{student_row['班號']} | ").font.size = Pt(9)

        # ── 缺席學生：只顯示缺席說明，不生成成績內容 ──
        if student_name in absent_set:
            run = info_para.add_run("【本次缺席】")
            run.font.size = Pt(9); run.font.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            abs_para = combined_doc.add_paragraph("此學生本次考試缺席，未有成績記錄。")
            abs_para.runs[0].font.size = Pt(9)
            abs_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            combined_doc.add_paragraph()
            continue

        # ── 出席學生：顯示正常成績 ──
        df_student = df.loc[student_name]
        student_total_score = float(df_student.sum())
        student_total_max = int(max_scores.sum())
        score_rate = student_total_score / student_total_max * 100

        run = info_para.add_run(f"總分：{student_total_score:.1f}/{student_total_max} ({score_rate:.1f}%)")
        run.font.size = Pt(9)
        run.font.bold = True

        combined_doc.add_paragraph()

        # 各題成績表
        table = combined_doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        # 表頭
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '題號'
        hdr_cells[1].text = '滿分'
        hdr_cells[2].text = '得分'
        hdr_cells[3].text = '得分率'

        for cell in hdr_cells:
            set_cell_background(cell, '4472C4')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True

        # 填充各題資料
        for question in df_student.index:
            max_s = max_scores[question]
            score = df_student[question]
            rate = score / max_s * 100 if max_s > 0 else 0

            row_cells = table.add_row().cells
            row_cells[0].text = str(question)
            row_cells[1].text = str(int(max_s))
            row_cells[2].text = str(int(score)) if score == int(score) else f"{score:.1f}"
            row_cells[3].text = f"{rate:.0f}%"

        # 分隔符（非最後一個學生）
        if idx < len(sorted_students) - 1:
            combined_doc.add_paragraph()
            combined_doc.add_page_break()

    # ========== 儲存 DOCX ==========
    combined_docx_path = os.path.join(output_dir, f"{exam_info['file_prefix']}_全班報告.docx")
    combined_doc.save(combined_docx_path)

    print(f"✅ 全班 DOCX 報告已生成：{combined_docx_path}")

    return combined_docx_path



def generate_all_reports(df, max_scores, item_df, exam_info, class_info, output_dir, pass_rate=0.4, absent_set=None):
    """為所有學生生成報告"""
    absent_set = absent_set or set()

    os.makedirs(output_dir, exist_ok=True)

    # 保存 class_info 供 PDF 合併使用
    df_class_info = class_info.copy()

    total_scores = df.sum(axis=1)
    total_max = int(max_scores.sum())

    # 計算全班平均分
    class_avg_scores = item_df["平均分"].sum()
    class_total_max = total_max

    report_data = []
    pdf_count = 0

    print(f"\n📄 生成個人報告（v2.4 增強版）...")

    # 按班別和班號排序
    df_sorted_idx = df.index.tolist()
    class_info_dict = dict(zip(class_info['中文姓名'], 
                               zip(class_info['班別'], class_info['班號'])))

    def get_sort_key(name):
        if name in class_info_dict:
            class_code, class_num = class_info_dict[name]
            # 班別排序：5A < 5B < 5C ... 6A < 6B ...
            # 班號排序：01 < 02 < ... < 99
            try:
                return (class_code, int(class_num))
            except:
                return (class_code, 999)
        return ("ZZ", 999)

    df_sorted_idx.sort(key=get_sort_key)

    absent_count = 0
    for idx, student_name in enumerate(df_sorted_idx, 1):
        student_score = total_scores[student_name]

        # 取得班別和班號
        if student_name in class_info_dict:
            class_code, class_num = class_info_dict[student_name]
        else:
            class_code, class_num = "N/A", "00"

        # ── 缺席學生：生成缺席通知頁 ──
        if student_name in absent_set:
            doc = Document()
            sec = doc.sections[0]
            sec.top_margin = sec.bottom_margin = Inches(1.0)
            sec.left_margin = sec.right_margin = Inches(1.2)
            doc.add_paragraph()
            doc.add_paragraph()
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_title.add_run("個人成績報告")
            r.font.name = "Microsoft JhengHei"; r.font.size = Pt(20)
            r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = p_name.add_run(student_name)
            rn.font.name = "Microsoft JhengHei"; rn.font.size = Pt(16)
            rn.font.color.rgb = RGBColor(0x1F, 0x6F, 0xB8)

            p_exam = doc.add_paragraph()
            p_exam.alignment = WD_ALIGN_PARAGRAPH.CENTER
            re_ = p_exam.add_run(exam_info.get("exam_title", ""))
            re_.font.name = "Microsoft JhengHei"; re_.font.size = Pt(11)
            re_.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

            for _ in range(3):
                doc.add_paragraph()

            p_abs = doc.add_paragraph()
            p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ra = p_abs.add_run("本次缺席")
            ra.font.name = "Microsoft JhengHei"; ra.font.size = Pt(36)
            ra.font.bold = True; ra.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

            p_note = doc.add_paragraph()
            p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn2 = p_note.add_run("此學生本次考試缺席，未有成績記錄。")
            rn2.font.name = "Microsoft JhengHei"; rn2.font.size = Pt(11)
            rn2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            absent_count += 1

        else:
            # ── 出席學生：生成正常成績報告 ──
            doc = create_personal_report_v2_4(
                student_name,
                student_score,
                total_max,
                df.loc[student_name],
                max_scores,
                item_df,
                exam_info,
                class_avg_scores,
                class_total_max,
                class_info,
                pass_rate
            )

        # 儲存 Word 檔（文件名加班別和班號）
        filename_prefix = f"{class_code}{int(class_num):02d}{student_name}"
        word_filename = f"{output_dir}/{filename_prefix}_個人報告.docx"
        doc.save(word_filename)

        # 轉換為 PDF
        pdf_filename = f"{output_dir}/{filename_prefix}_個人報告.pdf"
        if docx_to_pdf(word_filename, pdf_filename):
            pdf_count += 1

        report_data.append({
            "姓名":   student_name,
            "出席狀態": "缺席" if student_name in absent_set else "出席",
            "總分":   "缺席" if student_name in absent_set else int(student_score),
            "得分率":  "缺席" if student_name in absent_set else f"{student_score/total_max*100:.1f}%",
            "Word":  word_filename,
            "PDF":   pdf_filename if os.path.exists(pdf_filename) else "未生成",
        })

        if idx % 5 == 0:
            print(f"    已生成 {idx} 份報告...")

    present_count = len(report_data) - absent_count
    print(f"    共生成 {len(report_data)} 份個人報告（出席 {present_count} 人 / 缺席 {absent_count} 人）")
    print(f"    PDF 成功轉換：{pdf_count}/{len(report_data)} 份")

    # 嘗試合併 PDF
    merged_pdf_path = merge_individual_pdfs(output_dir, df_class_info, 
                                             exam_info.get('file_prefix', 'report'))

    return report_data, output_dir, merged_pdf_path
