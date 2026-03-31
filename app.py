"""
DSE 試卷分析系統 - 網頁版
Streamlit app (app.py)
"""
import streamlit as st
import pandas as pd
import numpy as np
import io, os, sys, zipfile, tempfile, re

# ── 載入分析模組（與 app.py 同資料夾）──
sys.path.insert(0, os.path.dirname(__file__))
import exam_item_analysis as ea
import individual_report  as ir

# ══════════════════════════════════════════════════════════════
# 頁面設定
# ══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="DSE 試卷分析系統",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── 全域 CSS ──
st.markdown("""
<style>
    .main-header{font-size:1.8rem;font-weight:700;color:#1F4788;margin-bottom:0.2rem;}
    .sub-header{font-size:1rem;color:#555;margin-bottom:1.5rem;}
    .step-badge{background:#1F4788;color:white;border-radius:50%;
                width:28px;height:28px;display:inline-flex;
                align-items:center;justify-content:center;
                font-weight:bold;margin-right:8px;}
    .success-box{background:#d4edda;border-left:4px solid #28a745;
                 padding:0.75rem 1rem;border-radius:4px;margin:0.5rem 0;}
    .warn-box{background:#fff3cd;border-left:4px solid #ffc107;
              padding:0.75rem 1rem;border-radius:4px;margin:0.5rem 0;}
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# 工具函數：load_data 的 BytesIO 版本
# ══════════════════════════════════════════════════════════════
def load_data_from_bytes(file_bytes: bytes):
    """從 bytes 讀取 scores.xlsx（適用 Streamlit UploadedFile）"""
    buf = io.BytesIO(file_bytes)
    df_raw    = pd.read_excel(buf, header=None)
    col_names = df_raw.iloc[2].tolist()

    row3_vals = df_raw.iloc[3].tolist()
    paper_labels = {"P1","P2","P3","P4"}
    row3_str  = [str(v).strip() for v in row3_vals if pd.notna(v) and str(v).strip()!=""]
    has_paper_row = any(v in paper_labels for v in row3_str)

    if has_paper_row:
        paper_row   = row3_vals
        max_row     = df_raw.iloc[4].tolist()
        student_raw = df_raw.iloc[5:].copy()
    else:
        paper_row   = None
        max_row     = row3_vals
        student_raw = df_raw.iloc[4:].copy()

    student_raw.columns = col_names
    student_raw = student_raw[
        student_raw["中文姓名"].notna() &
        ~student_raw["中文姓名"].astype(str).str.contains("說明|輸入", na=False)
    ]

    absent_set = set()
    absent_col_candidates = ["缺席","Absent","absent","缺考"]
    absent_col = next((c for c in col_names if str(c) in absent_col_candidates), None)
    if absent_col:
        for _, row in student_raw.iterrows():
            val = row.get(absent_col,"")
            if pd.notna(val) and str(val).strip() not in ["","0","nan"]:
                absent_set.add(str(row["中文姓名"]).strip())

    info_cols     = ["班別","班號","英文姓名","中文姓名"] + absent_col_candidates
    question_cols = [c for c in col_names
                     if isinstance(c,str) and c.strip()!="" and c not in info_cols]
    max_scores = pd.Series(max_row, index=col_names)[question_cols].astype(float)
    score_data = student_raw[question_cols].astype(float)
    score_data.index = student_raw["中文姓名"].values
    score_data.index.name = "姓名"
    score_data.fillna(0, inplace=True)

    # class_info
    ci = pd.DataFrame()
    ci["班別"]    = student_raw.iloc[:,0].values
    ci["班號"]    = student_raw.iloc[:,1].values
    ci["中文姓名"] = student_raw["中文姓名"].values
    ci = ci[ci["中文姓名"].notna()].reset_index(drop=True)

    paper_map = {}
    if has_paper_row:
        paper_series = pd.Series(paper_row, index=col_names)
        for q in question_cols:
            val = str(paper_series.get(q,"P1")).strip()
            paper_map[q] = val if val in paper_labels else "P1"
    else:
        paper_map = {q:"P1" for q in question_cols}

    return score_data, max_scores, absent_set, paper_map, ci


def export_excel_bytes(item_df, group_df, student_df, stats_df, exam_title):
    """匯出 Excel 到 BytesIO"""
    buf = io.BytesIO()
    from openpyxl.styles import (PatternFill, Font, Border, Side,
                                  Alignment, GradientFill)
    from openpyxl.utils import get_column_letter

    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin"))
    header_fill = PatternFill("solid", fgColor="1F4788")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    color_fills = {
        "🟢 容易": PatternFill("solid", fgColor="C6EFCE"),
        "🟡 適中": PatternFill("solid", fgColor="FFEB9C"),
        "🔴 困難": PatternFill("solid", fgColor="FFC7CE"),
        "⭐ 優良": PatternFill("solid", fgColor="C6EFCE"),
        "✅ 良好": PatternFill("solid", fgColor="DDEBF7"),
        "⚠️ 尚可": PatternFill("solid", fgColor="FFEB9C"),
        "❌ 不佳": PatternFill("solid", fgColor="FFC7CE"),
    }

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        item_df.to_excel(writer, sheet_name="1_試題分析", index=False)
        group_df.to_excel(writer, sheet_name="2_大題分析", index=False)

        student_df_sorted = student_df.copy()
        rank_col = pd.to_numeric(student_df_sorted["排名"], errors="coerce")
        student_df_sorted = student_df_sorted.iloc[rank_col.argsort(kind="stable")]
        absent_mask = student_df_sorted["出席狀態"] == "缺席"
        student_df_sorted = pd.concat([student_df_sorted[~absent_mask],
                                       student_df_sorted[absent_mask]])
        student_df_sorted.to_excel(writer, sheet_name="3_學生成績", index=True)
        stats_df.to_excel(writer, sheet_name="4_全班統計", index=False)

        wb = writer.book
        # 格式化試題分析
        ws = wb["1_試題分析"]
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center")
                for keyword, fill in color_fills.items():
                    if str(cell.value) == keyword:
                        cell.fill = fill
        for col in ws.columns:
            ws.column_dimensions[get_column_letter(col[0].column)].width = 14

    buf.seek(0)
    return buf.read()


def _docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes | None:
    """用 LibreOffice 將 docx bytes 轉為 PDF bytes（需 packages.txt: libreoffice）"""
    import subprocess, tempfile
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "report.docx")
        pdf_path  = os.path.join(tmpdir, "report.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        for lo_cmd in ["libreoffice", "libreoffice7.6", "libreoffice7.5",
                       "/usr/bin/libreoffice", "/usr/lib/libreoffice/program/soffice"]:
            try:
                result = subprocess.run(
                    [lo_cmd, "--headless", "--convert-to", "pdf",
                     "--outdir", tmpdir, docx_path],
                    capture_output=True, timeout=60
                )
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        return f.read()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
    return None


def _merge_pdf_bytes(pdf_bytes_list: list) -> bytes:
    """合併多個 PDF bytes 為一個 PDF"""
    from pypdf import PdfWriter
    writer = PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        buf = io.BytesIO(pdf_bytes)
        from pypdf import PdfReader
        reader = PdfReader(buf)
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out.read()


def generate_reports_zip(df, max_scores, item_df, exam_info, class_info,
                          pass_rate, absent_set, gen_pdf=True):
    """
    生成個人報告 ZIP（BytesIO）
    回傳：(docx_zip_bytes, pdf_zip_bytes, merged_pdf_bytes)
    pdf_zip_bytes / merged_pdf_bytes 在 LibreOffice 不可用時為 None
    """
    total_scores = df.sum(axis=1)
    total_max    = int(max_scores.sum())
    class_avg    = item_df["平均分"].sum()

    class_info_dict = {}
    if class_info is not None and len(class_info):
        class_info_dict = dict(zip(
            class_info["中文姓名"],
            zip(class_info["班別"].astype(str), class_info["班號"].astype(str))
        ))

    docx_entries = []   # [(fname_base, docx_bytes), ...]
    pdf_entries  = []   # [(fname_base, pdf_bytes), ...]

    for student_name in df.index:
        student_score = total_scores[student_name]
        class_code, class_num = class_info_dict.get(student_name, ("", "00"))
        try:
            class_num_int = int(float(class_num))
        except:
            class_num_int = 0

        if student_name in absent_set:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
            doc = Document()
            sec = doc.sections[0]
            sec.top_margin = sec.bottom_margin = Inches(1.0)
            sec.left_margin = sec.right_margin = Inches(1.2)
            for _ in range(3): doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("本次缺席")
            r.font.size = Pt(36); r.font.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            doc = ir.create_personal_report_v2_4(
                student_name, student_score, total_max,
                df.loc[student_name], max_scores, item_df,
                exam_info, class_avg, total_max,
                class_info, pass_rate
            )

        fname_base = f"{class_code}{class_num_int:02d}{student_name}_個人報告"
        doc_buf = io.BytesIO()
        doc.save(doc_buf)
        docx_bytes = doc_buf.getvalue()
        docx_entries.append((fname_base, docx_bytes))

        if gen_pdf:
            pdf_bytes = _docx_bytes_to_pdf_bytes(docx_bytes)
            if pdf_bytes:
                pdf_entries.append((fname_base, pdf_bytes))

    # 打包 Word ZIP
    docx_buf = io.BytesIO()
    with zipfile.ZipFile(docx_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname_base, docx_bytes in docx_entries:
            zf.writestr(f"{fname_base}.docx", docx_bytes)
    docx_buf.seek(0)
    docx_zip = docx_buf.read()

    # 打包 PDF ZIP + 合併 PDF
    pdf_zip = merged_pdf = None
    if pdf_entries:
        pdf_buf = io.BytesIO()
        with zipfile.ZipFile(pdf_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname_base, pdf_bytes in pdf_entries:
                zf.writestr(f"{fname_base}.pdf", pdf_bytes)
        pdf_buf.seek(0)
        pdf_zip = pdf_buf.read()
        try:
            merged_pdf = _merge_pdf_bytes([pb for _, pb in pdf_entries])
        except Exception:
            merged_pdf = None

    return docx_zip, pdf_zip, merged_pdf


# ══════════════════════════════════════════════════════════════
# 追蹤分析工具函數
# ══════════════════════════════════════════════════════════════
def export_tracking_excel_bytes(pct_matrix, rank_matrix, student_info,
                                 class_stats, exam_labels, pass_rate, subject):
    """追蹤報告 Excel → BytesIO"""
    buf = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = os.path.join(tmpdir, "tracking.xlsx")
        import performance_tracker as pt
        pt.export_tracking_excel(
            pct_matrix, rank_matrix, student_info, class_stats,
            exam_labels, tmp_path, pass_rate, subject
        )
        with open(tmp_path, "rb") as f:
            buf.write(f.read())
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# 側邊欄導航
# ══════════════════════════════════════════════════════════════
st.sidebar.markdown("## 📊 DSE 試卷分析系統")
st.sidebar.markdown("---")
page = st.sidebar.radio(
    "功能選擇",
    ["📋 試卷分析", "📈 成績追蹤"],
    label_visibility="collapsed"
)
st.sidebar.markdown("---")



# ══════════════════════════════════════════════════════════════
# 頁面一：試卷分析
# ══════════════════════════════════════════════════════════════
if page == "📋 試卷分析":
    st.markdown('''<div class="main-header">📋 試卷分析</div>
<div class="sub-header">上載成績表，即時生成試題分析、學生報告及圖表</div>''',
                unsafe_allow_html=True)

    # ── Step 1: 考試資訊 ──
    st.markdown("### ① 考試資訊")
    col1, col2, col3 = st.columns(3)
    with col1:
        year_input = st.text_input("年度（如 2526）", value="2526", max_chars=4)
        exam_type  = st.selectbox("考試類別", ["上學期測驗","上學期考試","下學期測驗","下學期考試"])
    with col2:
        form_label  = st.selectbox("年級", ["F1","F2","F3","F4","F5","F6"])
        subject_label = st.text_input("科目（如 BAFS）", value="BAFS")
    with col3:
        pass_rate_pct = st.selectbox("及格線", ["40%（高中）","50%（初中）"])
        pass_rate = 0.4 if "40" in pass_rate_pct else 0.5

    # 試卷設定
    st.markdown("### ② 試卷結構")
    num_papers = st.selectbox("試卷數目", [1,2,3,4],
                               format_func=lambda x: f"{x} 份試卷")
    paper_weights = {"P1": 1.0}
    if num_papers > 1:
        st.info("各試卷比例合計必須為 100%")
        cols = st.columns(num_papers)
        tmp_weights = {}
        for i, col in enumerate(cols, 1):
            with col:
                w = col.number_input(f"卷{i}（P{i}）比例 %",
                                      min_value=1, max_value=99,
                                      value=100//num_papers, step=1,
                                      key=f"pw_{i}")
                tmp_weights[f"P{i}"] = w
        total_w = sum(tmp_weights.values())
        if total_w != 100:
            st.warning(f"⚠️ 目前合計：{total_w}%，必須等於 100%")
        else:
            paper_weights = {k: v/100 for k, v in tmp_weights.items()}
            st.success(f"✅ 各卷比例設定正確")

    # 組合 exam_info
    EXAM_CODES = {"上學期測驗":"T1T","上學期考試":"T1E","下學期測驗":"T2T","下學期考試":"T2E"}
    year_label  = f"20{year_input[:2]}-20{year_input[2:]}" if len(year_input)==4 else year_input
    exam_title  = f"{year_label} {exam_type}｜{form_label} {subject_label}"
    file_prefix = f"{year_input}_{EXAM_CODES.get(exam_type,'EXAM')}_{form_label}_{subject_label}"
    exam_info = {
        "year_label": year_label, "exam_type_label": exam_type,
        "exam_type_code": EXAM_CODES.get(exam_type,"EXAM"),
        "subject_label": subject_label, "form_label": form_label,
        "exam_title": exam_title, "file_prefix": file_prefix,
        "pass_rate": pass_rate, "paper_weights": paper_weights,
        "num_papers": num_papers,
    }

    # ── Step 3: 上載檔案 ──
    st.markdown("### ③ 上載成績表")
    uploaded = st.file_uploader("選擇 scores.xlsx", type=["xlsx"],
                                  help="第3行為題號，第4行為【試卷】（多試卷時），之後為【滿分】，再下為學生成績")

    if uploaded:
        raw_bytes = uploaded.read()
        try:
            df, max_scores, absent_set, paper_map, class_info = load_data_from_bytes(raw_bytes)
            papers_in_excel = sorted(set(paper_map.values()))
            if len(papers_in_excel) > 1 and num_papers == 1:
                st.warning(f"⚠️ 偵測到試卷行 {papers_in_excel}，但設定為單試卷，請調整上方試卷數目")

            st.markdown(f'''<div class="success-box">
            ✅ 成功載入：<b>{len(df)} 名學生</b>，<b>{len(df.columns)} 道題目</b>
            {"｜缺席：" + "、".join(sorted(absent_set)) if absent_set else ""}
            </div>''', unsafe_allow_html=True)

            # 預覽
            with st.expander("📋 預覽學生名單", expanded=False):
                st.dataframe(class_info, use_container_width=True)

            # ── 分析按鈕 ──
            st.markdown("### ④ 開始分析")

            # ── 分析按鈕（結果存入 session_state，避免重複分析）──
            col_btn, col_reset = st.columns([3, 1])
            with col_btn:
                analyze_btn = st.button("🚀 開始分析", type="primary", use_container_width=True)
            with col_reset:
                if st.button("🔄 重新上載", use_container_width=True):
                    for k in ["analysis_done","item_df","student_df","stats_df",
                              "group_df","excel_bytes","docx_zip","pdf_zip",
                              "merged_pdf","charts_png_zip","scores_num",
                              "item_plot","fig_data"]:
                        st.session_state.pop(k, None)
                    st.rerun()

            if analyze_btn:
                with st.spinner("分析中，請稍候..."):
                    weighted_scores, paper_pct, _ = ea.calc_weighted_scores(
                        df, max_scores, paper_weights, paper_map)
                    item_df    = ea.item_analysis(df.copy(), max_scores, absent_set)
                    student_df, stats_df = ea.student_summary(
                        df.copy(), max_scores, pass_rate, absent_set,
                        paper_weights=paper_weights, paper_pct=paper_pct,
                        weighted_scores=weighted_scores, num_papers=num_papers)
                    group_df   = ea.question_group_analysis(df.copy(), max_scores, item_df)

                with st.spinner("生成報告檔案中..."):
                    excel_bytes = export_excel_bytes(
                        item_df, group_df, student_df, stats_df, exam_title)
                    docx_zip, pdf_zip, merged_pdf = generate_reports_zip(
                        df, max_scores, item_df, exam_info,
                        class_info, pass_rate, absent_set, gen_pdf=True)

                # 計算圖表資料（供 Plotly 互動顯示用）
                item_plot = item_df.copy()
                item_plot["得分率 %"] = (item_plot["平均分"] / item_plot["滿分"] * 100).round(1)
                scores_num = pd.to_numeric(student_df.get(
                    "總分(加權)" if "總分(加權)" in student_df.columns else "總分",
                    pd.Series()), errors="coerce").dropna()

                # ── 用原版 create_charts 生成完全一致的圖表（含熱力圖）──
                with st.spinner("生成圖表中（使用 kaleido==0.1.0.post1）..."):
                    try:
                        charts_png = ea.create_charts(
                            df.copy(), max_scores, item_df.copy(), student_df.copy(),
                            exam_title, chart_dir=None, absent_set=absent_set,
                            return_bytes=True
                        ) or {}
                        if charts_png:
                            st.success(f"✅ 成功生成 {len(charts_png)} 張圖表")
                        else:
                            st.warning("⚠️ 圖表生成為空，請確認 kaleido==0.1.0.post1 已安裝")
                    except Exception as _ce:
                        st.error(f"❌ 圖表生成失敗：{_ce}")
                        st.info("請確認 requirements.txt 中含 kaleido==0.1.0.post1 並重新部署")
                        charts_png = {}

                charts_zip_buf = io.BytesIO()
                with zipfile.ZipFile(charts_zip_buf, "w", zipfile.ZIP_DEFLATED) as czf:
                    labels = {
                        "01_difficulty_discrimination.png": "01_難度鑑別度分佈.png",
                        "02_score_rate_by_question.png":    "02_各題得分率排行.png",
                        "03_student_score_distribution.png":"03_全班總分分佈.png",
                        "04_class_heatmap.png":             "04_全班答題熱力圖.png",
                    }
                    for fname, png in charts_png.items():
                        czf.writestr(labels.get(fname, fname), png)
                charts_zip_buf.seek(0)

                # 存入 session_state
                ss = st.session_state
                ss.analysis_done  = True
                ss.item_df        = item_df
                ss.student_df     = student_df
                ss.stats_df       = stats_df
                ss.group_df       = group_df
                ss.excel_bytes    = excel_bytes
                ss.docx_zip       = docx_zip
                ss.pdf_zip        = pdf_zip
                ss.merged_pdf     = merged_pdf
                ss.charts_png     = charts_png
                ss.charts_png_zip = charts_zip_buf.read()
                ss.item_plot      = item_plot
                ss.scores_num     = scores_num
                ss.file_prefix    = file_prefix
                ss.exam_title     = exam_title
                st.success("✅ 分析完成！所有報告已就緒，可直接下載。")

            # ── 顯示結果（從 session_state 讀取，不重新分析）──
            if st.session_state.get("analysis_done"):
                ss         = st.session_state
                item_df    = ss.item_df
                student_df = ss.student_df
                stats_df   = ss.stats_df
                group_df   = ss.group_df
                item_plot  = ss.item_plot
                scores_num = ss.scores_num
                fp         = ss.file_prefix

                tab1, tab2, tab3, tab4 = st.tabs(
                    ["📊 試題分析","📋 大題分析","👨‍🎓 學生成績","📈 全班統計"])
                with tab1:
                    st.dataframe(item_df, use_container_width=True, height=400)
                with tab2:
                    st.dataframe(group_df, use_container_width=True)
                with tab3:
                    st.dataframe(student_df, use_container_width=True, height=400)
                with tab4:
                    st.dataframe(stats_df, use_container_width=True)

                # ── 圖表預覽（直接顯示原版 PNG，與下載完全一致）──
                st.markdown("### 📊 圖表預覽")
                charts_png_ss = ss.charts_png
                chart_order = [
                    ("01_difficulty_discrimination.png", "① 難度－鑑別度分佈"),
                    ("02_score_rate_by_question.png",    "② 各題得分率排行"),
                    ("03_student_score_distribution.png","③ 全班總分分佈"),
                    ("04_class_heatmap.png",             "④ 全班答題熱力圖"),
                ]
                for fname, label in chart_order:
                    if fname in charts_png_ss:
                        st.markdown(f"**{label}**")
                        st.image(charts_png_ss[fname], use_container_width=True)

                # ── 下載區（全部從 session_state 讀取，無需重新分析）──
                st.markdown("### ⬇️ 下載報告")

                # 一鍵下載全部
                all_zip_buf = io.BytesIO()
                with zipfile.ZipFile(all_zip_buf, "w", zipfile.ZIP_DEFLATED) as azf:
                    azf.writestr(f"{fp}_analysis.xlsx",        ss.excel_bytes)
                    azf.writestr(f"{fp}_個人報告_Word.zip",     ss.docx_zip)
                    if ss.pdf_zip:
                        azf.writestr(f"{fp}_個人報告_PDF.zip", ss.pdf_zip)
                    if ss.merged_pdf:
                        azf.writestr(f"{fp}_全班個人報告.pdf",  ss.merged_pdf)
                    azf.writestr(f"{fp}_圖表.zip",             ss.charts_png_zip)
                all_zip_buf.seek(0)
                st.download_button(
                    "📦 一鍵下載全部檔案 ZIP",
                    data=all_zip_buf.read(),
                    file_name=f"{fp}_全部報告.zip",
                    mime="application/zip",
                    use_container_width=True,
                    type="primary")

                st.markdown("---")

                # 個別下載
                dl1, dl2, dl3 = st.columns(3)
                with dl1:
                    st.download_button(
                        "📥 Excel 分析報告",
                        data=ss.excel_bytes,
                        file_name=f"{fp}_analysis.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True)
                with dl2:
                    st.download_button(
                        "📥 個人報告 Word ZIP",
                        data=ss.docx_zip,
                        file_name=f"{fp}_個人報告_Word.zip",
                        mime="application/zip",
                        use_container_width=True)
                with dl3:
                    st.download_button(
                        "📥 圖表 ZIP（PNG）",
                        data=ss.charts_png_zip,
                        file_name=f"{fp}_圖表.zip",
                        mime="application/zip",
                        use_container_width=True)

                dl4, dl5 = st.columns(2)
                with dl4:
                    if ss.pdf_zip:
                        st.download_button(
                            "📥 個人報告 PDF ZIP",
                            data=ss.pdf_zip,
                            file_name=f"{fp}_個人報告_PDF.zip",
                            mime="application/zip",
                            use_container_width=True)
                    else:
                        st.info("📄 PDF：需 packages.txt 安裝 LibreOffice")
                with dl5:
                    if ss.merged_pdf:
                        st.download_button(
                            "📥 合併個人報告 PDF",
                            data=ss.merged_pdf,
                            file_name=f"{fp}_全班個人報告.pdf",
                            mime="application/pdf",
                            use_container_width=True)
                    else:
                        st.info("📄 合併 PDF：需 LibreOffice + pypdf")


        except Exception as e:
            st.error(f"❌ 載入失敗：{e}")
            import traceback
            st.code(traceback.format_exc())


# ══════════════════════════════════════════════════════════════
# 頁面二：成績追蹤
# ══════════════════════════════════════════════════════════════
elif page == "📈 成績追蹤":
    st.markdown('''<div class="main-header">📈 成績追蹤</div>
<div class="sub-header">上載多次考試的分析報告，生成跨試成績追蹤</div>''',
                unsafe_allow_html=True)

    import performance_tracker as pt

    # ── Step 1: 設定 ──
    st.markdown("### ① 追蹤設定")
    col1, col2 = st.columns(2)
    with col1:
        track_subject  = st.text_input("科目名稱", value="BAFS")
        track_pass_pct = st.selectbox("及格線", ["40%（高中）","50%（初中）"], key="tp")
        track_pass_rate = 0.4 if "40" in track_pass_pct else 0.5
    with col2:
        st.markdown("**分析檔案命名規則**")
        st.code("2526_T1E_F5_BAFS_analysis.xlsx\n2526_T2E_F5_BAFS_analysis.xlsx", language=None)
        st.caption("格式：年度_類別_年級_科目_analysis.xlsx")

    # ── Step 2: 上載 ──
    st.markdown("### ② 上載分析報告（可多選）")
    track_files = st.file_uploader(
        "選擇 *_analysis.xlsx 檔案（可多選）",
        type=["xlsx"], accept_multiple_files=True,
        help="即之前試卷分析下載的 Excel 報告")

    if track_files:
        st.markdown(f"已上載 **{len(track_files)}** 個檔案：")
        exam_files_data = []
        for f in track_files:
            year, etype, form, subject = pt.parse_filename(f.name)
            tag = " ".join(filter(None,[form,subject]))
            st.caption(f"· {f.name}  [{tag or '舊格式'}]  {year} {etype}")
            if year and etype:
                exam_files_data.append({
                    "file_obj": f, "file": f.name,
                    "year": year, "type": etype,
                    "form": form, "subject": subject,
                })
            else:
                st.warning(f"⚠️ {f.name} 無法解析檔名，請確認命名格式")

        # 過濾選項
        forms_found    = sorted(set(ef["form"]    for ef in exam_files_data if ef["form"]))
        subjects_found = sorted(set(ef["subject"] for ef in exam_files_data if ef["subject"]))

        col1, col2 = st.columns(2)
        with col1:
            form_filter = st.selectbox(
                "年級過濾", ["（全部）"] + forms_found) if len(forms_found)>1 else (
                forms_found[0] if forms_found else "（全部）")
        with col2:
            subj_filter = st.selectbox(
                "科目過濾", ["（全部）"] + subjects_found) if len(subjects_found)>1 else (
                subjects_found[0] if subjects_found else "（全部）")

        filtered = [ef for ef in exam_files_data
                    if (form_filter in ["（全部）",""] or ef["form"]==form_filter)
                    and (subj_filter in ["（全部）",""] or ef["subject"]==subj_filter)]
        st.info(f"過濾後：{len(filtered)} 個檔案納入分析")

        # ── Step 3: 班別班號（選填）──
        st.markdown("### ③ 上載最新 scores.xlsx（選填，補充班別班號）")
        ci_file = st.file_uploader("選擇 scores.xlsx（選填）", type=["xlsx"], key="ci_upload")
        class_info_df = None
        if ci_file:
            try:
                raw = pd.read_excel(io.BytesIO(ci_file.read()), header=None)
                ci = pd.DataFrame()
                ci["班別"]    = raw.iloc[4:,0].values
                ci["班號"]    = raw.iloc[4:,1].values
                ci["中文姓名"] = raw.iloc[4:,3].values
                ci = ci[ci["中文姓名"].notna() &
                        ~ci["中文姓名"].astype(str).str.contains("說明|輸入",na=False)]
                class_info_df = ci.reset_index(drop=True)
                st.success(f"✅ 讀取 {len(ci)} 位學生班別資訊")
            except Exception as e:
                st.warning(f"⚠️ 讀取 scores.xlsx 失敗：{e}")

        # ── Step 4: 生成 ──
        if st.button("🚀 生成追蹤報告", type="primary",
                     use_container_width=True, disabled=len(filtered)==0):
            with st.spinner("建立成績矩陣中..."):
                # 將上載檔案寫入暫存目錄
                with tempfile.TemporaryDirectory() as tmpdir:
                    ef_list = []
                    for ef in filtered:
                        tmp_path = os.path.join(tmpdir, ef["file"])
                        ef["file_obj"].seek(0)
                        with open(tmp_path, "wb") as out:
                            out.write(ef["file_obj"].read())
                        ef_list.append({
                            "file": tmp_path, "year": ef["year"],
                            "type": ef["type"], "form": ef["form"],
                            "subject": ef["subject"],
                        })

                    pct_matrix, rank_matrix, exam_labels, student_info = (
                        pt.build_tracking_matrix(ef_list, class_info_df))

                    if pct_matrix is None:
                        st.error("❌ 無法建立成績矩陣，請確認檔案格式")
                    else:
                        class_stats = pt.calc_class_stats(pct_matrix, track_pass_rate)
                        st.success(f"✅ 矩陣：{pct_matrix.shape[0]} 位學生 × {pct_matrix.shape[1]} 次考試")

                        # 顯示結果
                        tab1, tab2 = st.tabs(["📊 全班趨勢","👨‍🎓 學生成績矩陣"])
                        with tab1:
                            import plotly.graph_objects as go
                            fig = go.Figure()
                            for col in class_stats.columns[1:]:
                                vals = pd.to_numeric(class_stats[col], errors="coerce")
                                if not vals.isna().all():
                                    fig.add_trace(go.Scatter(
                                        x=class_stats["考試"],
                                        y=vals, mode="lines+markers",
                                        name=col))
                            fig.update_layout(title="全班成績趨勢",
                                              xaxis_title="考試",
                                              yaxis_title="分數")
                            st.plotly_chart(fig, use_container_width=True)
                        with tab2:
                            st.dataframe(pct_matrix.round(1), use_container_width=True, height=400)

                        # ── 生成全部報告 ──
                        st.markdown("### ⬇️ 下載追蹤報告")
                        years    = sorted(set(ef["year"] for ef in filtered))
                        fp_track = f"{'_'.join(years)}_{track_subject}"

                        # Excel
                        tracking_excel = None
                        try:
                            tracking_excel = export_tracking_excel_bytes(
                                pct_matrix, rank_matrix, student_info,
                                class_stats, exam_labels,
                                track_pass_rate, track_subject)
                        except Exception as e:
                            st.error(f"❌ 生成 Excel 失敗：{e}")
                            import traceback; st.code(traceback.format_exc())

                        # Word + PDF
                        tracking_docx = tracking_pdf = None
                        with st.spinner("生成 Word 追蹤報告（含個人趨勢圖）..."):
                            try:
                                tracking_docx, tracking_pdf = pt.generate_tracking_report_bytes(
                                    pct_matrix, rank_matrix, student_info,
                                    class_stats, exam_labels,
                                    fp_track, track_subject, track_pass_rate)
                            except Exception as e:
                                st.warning(f"⚠️ Word 報告生成失敗：{e}")

                        # 一鍵下載全部
                        if tracking_excel or tracking_docx:
                            all_track_buf = io.BytesIO()
                            with zipfile.ZipFile(all_track_buf, "w", zipfile.ZIP_DEFLATED) as azf:
                                if tracking_excel:
                                    azf.writestr(f"{fp_track}_成績追蹤.xlsx", tracking_excel)
                                if tracking_docx:
                                    azf.writestr(f"{fp_track}_成績追蹤報告.docx", tracking_docx)
                                if tracking_pdf:
                                    azf.writestr(f"{fp_track}_成績追蹤報告.pdf", tracking_pdf)
                            all_track_buf.seek(0)
                            st.download_button(
                                "📦 一鍵下載全部追蹤報告 ZIP",
                                data=all_track_buf.read(),
                                file_name=f"{fp_track}_追蹤報告全部.zip",
                                mime="application/zip",
                                use_container_width=True,
                                type="primary")
                            st.markdown("---")

                        # 個別下載
                        tdl1, tdl2, tdl3 = st.columns(3)
                        with tdl1:
                            if tracking_excel:
                                st.download_button(
                                    "📥 成績追蹤 Excel",
                                    data=tracking_excel,
                                    file_name=f"{fp_track}_成績追蹤.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True)
                        with tdl2:
                            if tracking_docx:
                                st.download_button(
                                    "📥 追蹤報告 Word",
                                    data=tracking_docx,
                                    file_name=f"{fp_track}_成績追蹤報告.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True)
                            else:
                                st.info("Word 報告生成失敗")
                        with tdl3:
                            if tracking_pdf:
                                st.download_button(
                                    "📥 追蹤報告 PDF",
                                    data=tracking_pdf,
                                    file_name=f"{fp_track}_成績追蹤報告.pdf",
                                    mime="application/pdf",
                                    use_container_width=True)
                            else:
                                st.info("PDF 需 LibreOffice（packages.txt）")
